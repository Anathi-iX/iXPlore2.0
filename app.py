from flask import Flask, render_template, request, redirect, session, url_for, send_from_directory, jsonify, Response
from werkzeug.security import generate_password_hash, check_password_hash
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime, date
import smtplib
from email.mime.text import MIMEText
import msal
import os
import openai
from openai import OpenAI
import requests
from PyPDF2 import PdfReader, PdfMerger
from pdf2image import convert_from_path
import pytesseract
from dotenv import load_dotenv
from flask_migrate import Migrate
from extensions import db
import pytz
import tiktoken
import time
from werkzeug.utils import secure_filename
import fitz
import uuid
import mammoth
from docx import Document
from docx import Document as DocxDocument
import re, json, ast
import time, random
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
from docx.oxml.shared import qn
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

from PIL import Image, ImageOps, ImageFilter, ImageEnhance
import numpy as np
import io
import re
import pytesseract
from flask import send_file

import threading
from threading import Lock

import matplotlib
matplotlib.use("Agg")  # headless backend (important for servers)
import matplotlib.pyplot as plt
from textwrap import fill
from matplotlib.backends.backend_pdf import PdfPages

import math
from typing import Optional, List, Dict, Any, Tuple
import hashlib
import pickle
import threading

import os
import re
import tiktoken

load_dotenv()  # Load environment variables from .env file

# ---- Tesseract location (Windows) ----
tcmd = os.getenv("TESSERACT_CMD", "").strip()
if tcmd:
    pytesseract.pytesseract.tesseract_cmd = tcmd
    print("[TESSERACT] cmd =", pytesseract.pytesseract.tesseract_cmd)
    try:
        v = pytesseract.get_tesseract_version()
        print("[TESSERACT] version =", v)
    except Exception as e:
        print("[TESSERACT] ERROR:", e)

# Microsoft 365 / Azure AD configuration
client_id = os.getenv('MICROSOFT_CLIENT_ID')
tenant_id = os.getenv('MICROSOFT_TENANT_ID')
client_secret = os.getenv('MICROSOFT_CLIENT_SECRET')
redirect_uri = os.getenv('MICROSOFT_REDIRECT_URI')
authority = f"https://login.microsoftonline.com/{tenant_id}"

# MSAL app instance
msal_app = msal.ConfidentialClientApplication(client_id, authority=authority, client_credential=client_secret)

# Scopes required for accessing user's profile
SCOPES = ["User.Read"]

pdf_texts = {}

OCR_JOBS = {}          # job_id -> job dict
OCR_JOBS_LOCK = Lock() # guard concurrent access
FILE_OCR_JOB = {}      # filename -> job_id (so /ask can block until OCR done)

# Initialize OpenAI API Key
openai.api_key = os.getenv('OPENAI_API_KEY')

# === DEBUG SWITCHES ===
DEBUG_NO_OPENAI = False  # set to False when you are ready to spend money again

_CODEFENCE_RE = re.compile(r"^```(?:json)?\s*|\s*```$", re.IGNORECASE | re.MULTILINE)

# Configure SQL Alchemy
app = Flask(__name__)
app.secret_key = "your_secrete_key"
app.config['UPLOAD_FOLDER'] = 'uploads'

os.makedirs("uploads", exist_ok=True)
os.makedirs("temp_data", exist_ok=True)
os.makedirs("examples", exist_ok=True)  # ensure examples dir exists (for template)

app.config["SQLALCHEMY_DATABASE_URI"] = os.getenv('DATABASE_URL')
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False

db = SQLAlchemy(app)
migrate = Migrate(app, db)

south_africa_tz = pytz.timezone('Africa/Johannesburg')

# Database Model
class Feedback(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(120), nullable=False)
    feedback_text = db.Column(db.String(500), nullable=False)
    rating = db.Column(db.Integer, nullable=False)
    timestamp = db.Column(db.DateTime, nullable=False)

    def __repr__(self):
        return f"<Feedback {self.id}>"

###############################################################################

class Conversation(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_email = db.Column(db.String(255), nullable=False)
    filename = db.Column(db.String(255), nullable=False)
    question = db.Column(db.Text, nullable=False)
    answer = db.Column(db.Text, nullable=False)
    timestamp = db.Column(db.DateTime, default=datetime.utcnow)

    def __repr__(self):
        return f'<Conversation {self.id}>'

print(app.url_map)  # This will print all the routes defined in your app

# ---------------------- COMPANY TEMPLATE PATH ----------------------
COMPANY_TEMPLATE_PATH = os.path.join("examples", "QRS-TPL-010-B Internal Document_AI.docx")
EXAMPLE_TENDER_PATH = os.path.join("examples", "example_tender.pdf")
EXAMPLE_OUTPUT_PATH = os.path.join("examples", "example_output_methodology.docx")

# Sanity check at startup
if not os.path.exists(COMPANY_TEMPLATE_PATH):
    print(f"[BOOT] Company template not found: {COMPANY_TEMPLATE_PATH}")

EXAMPLE_TENDER_PATH = os.path.join("examples", "example_tender.pdf")
EXAMPLE_OUTPUT_PATH = os.path.join("examples", "example_output_methodology.docx")

_EMBED_MODEL = os.getenv("OPENAI_EMBED_MODEL", "text-embedding-3-small")
_EMBED_CACHE_DIR = os.path.join("temp_data", "embed_cache")
os.makedirs(_EMBED_CACHE_DIR, exist_ok=True)


def _cosine(a: List[float], b: List[float]) -> float:
    dot = 0.0
    na = 0.0
    nb = 0.0
    for x, y in zip(a, b):
        dot += x * y
        na += x * x
        nb += y * y
    if na <= 0 or nb <= 0:
        return 0.0
    return dot / math.sqrt(na * nb)


_PAGE_SPLIT_RE = re.compile(r"^===\s*PAGE\s+(\d+).*?===", re.IGNORECASE | re.MULTILINE)


def split_text_into_pages(extracted_text: str) -> List[Dict[str, Any]]:
    """
    Returns: [{"page": 1, "text": "..."} ...]
    Uses your markers: === PAGE n === or === PAGE n (embedded/ocr) ===
    Falls back to one chunk if markers are missing.
    """
    t = extracted_text or ""
    matches = list(_PAGE_SPLIT_RE.finditer(t))
    if not matches:
        return [{"page": 1, "text": t.strip()}]

    pages = []
    for i, m in enumerate(matches):
        page_no = int(m.group(1))
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(t)
        page_text = t[start:end].strip()
        if page_text:
            pages.append({"page": page_no, "text": page_text})
    return pages or [{"page": 1, "text": t.strip()}]


def _embed_texts(texts: List[str]) -> List[List[float]]:
    """
    Calls OpenAI embeddings (batched).
    """
    client = _get_openai_client()
    resp = client.embeddings.create(model=_EMBED_MODEL, input=texts)
    return [d.embedding for d in resp.data]


def _cache_path_for_doc(doc_key: str) -> str:
    safe = hashlib.sha256(doc_key.encode("utf-8")).hexdigest()[:24]
    return os.path.join(_EMBED_CACHE_DIR, f"{safe}.pkl")


def load_or_build_page_embeddings(doc_id: str, pages: List[Dict[str, Any]]) -> Dict[str, Any]:
    """
    Build once per extracted text (or per filename+size), store on disk.
    """
    cache_path = _cache_path_for_doc(doc_id)

    if os.path.exists(cache_path):
        try:
            with open(cache_path, "rb") as f:
                cached = pickle.load(f)
            # basic sanity check
            if cached and cached.get("count") == len(pages):
                return cached
        except Exception:
            pass

    # embed page texts (trim each page a bit to reduce cost, still works well)
    page_texts = []
    for p in pages:
        txt = p.get("text", "")
        page_texts.append(txt[:6000])

    embeddings = _embed_texts(page_texts)

    cached = {
        "count": len(pages),
        "pages": [{"page": pages[i]["page"], "text": pages[i]["text"]} for i in range(len(pages))],
        "embeddings": embeddings,
    }

    try:
        with open(cache_path, "wb") as f:
            pickle.dump(cached, f)
    except Exception:
        pass

    return cached

def is_summary_request(q: str) -> bool:
    ql = (q or "").lower()
    return any(k in ql for k in [
        "summarize", "summary", "summarise",
        "give me an overview", "overview",
        "key points", "main points", "high level",
        "tl;dr", "tldr"
    ])

def count_tokens(text: str, model: str) -> int:
    enc = tiktoken.encoding_for_model(model)
    return len(enc.encode(text or ""))

def chunk_by_tokens(text: str, model: str, max_tokens: int):
    enc = tiktoken.encoding_for_model(model)
    tokens = enc.encode(text or "")
    for i in range(0, len(tokens), max_tokens):
        yield enc.decode(tokens[i:i+max_tokens])

def answer_general_or_summary(question: str, full_text: str) -> str:
    model = os.getenv("OPENAI_QA_MODEL", "gpt-4o-mini")

    # Whole-doc summary (map-reduce)
    if is_summary_request(question):
        return summarize_whole_doc(full_text, model=model)

    # Old behavior for normal questions:
    # "Based on the following context, answer the question: ..."
    # We'll chunk to avoid token overflow, then stitch.
    max_ctx_tokens = 90000  # adjust to your model limits + safety
    chunks = list(chunk_by_tokens(full_text, model=model, max_tokens=max_ctx_tokens))
    if not chunks:
        return "No text was extracted from this document."

    answers = []
    for idx, chunk in enumerate(chunks, start=1):
        prompt = (
            "Based on the following context, answer the question:\n\n"
            f"{chunk}\n\n"
            f"Question: {question}\n\n"
            "If the answer is not in the context, say: Not found in extracted text."
        )
        text, _ = call_openai_text(model, prompt, temperature=0.0, max_output_tokens=450)
        answers.append((text or "").strip())

    # If doc is huge and multiple chunk answers were produced, combine them cleanly
    if len(answers) == 1:
        return answers[0]

    combine_prompt = (
        "Combine the following partial answers into ONE final answer.\n"
        "Remove duplicates. If answers conflict, say what differs.\n\n"
        + "\n\n---\n\n".join([f"PART {i+1}:\n{a}" for i, a in enumerate(answers)])
        + f"\n\nOriginal question: {question}"
    )
    final, _ = call_openai_text(model, combine_prompt, temperature=0.0, max_output_tokens=550)
    return (final or "").strip()

def summarize_whole_doc(full_text: str, model: str = "gpt-4o-mini") -> str:
    full_text = full_text or ""
    if not full_text.strip():
        return "No text was extracted from this document, so I can’t summarize it."

    # Pass 1: chunk summaries
    max_ctx_tokens = 60000
    chunks = list(chunk_by_tokens(full_text, model=model, max_tokens=max_ctx_tokens))

    partials = []
    for i, chunk in enumerate(chunks, start=1):
        prompt = (
            "You are summarizing a document. Write a concise, accurate summary of THIS PART.\n"
            "Return:\n"
            "- 5–12 bullet key points\n"
            "- Any dates, deadlines, amounts, and names (if present)\n"
            "- A 1-sentence 'what this part is about'\n\n"
            f"DOCUMENT PART {i}/{len(chunks)}:\n{chunk}"
        )
        part, _ = call_openai_text(model, prompt, temperature=0.2, max_output_tokens=700)
        partials.append((part or "").strip())

    # Pass 2: final synthesis (true whole-doc)
    final_prompt = (
        "Create a WHOLE-DOCUMENT summary from the chunk summaries below.\n"
        "Output format:\n"
        "1) One-paragraph overview\n"
        "2) 8–15 bullet key points\n"
        "3) Important dates/deadlines (if any)\n"
        "4) Action items / requirements (if any)\n\n"
        + "\n\n---\n\n".join([f"CHUNK SUMMARY {i+1}:\n{s}" for i, s in enumerate(partials)])
    )
    final, _ = call_openai_text(model, final_prompt, temperature=0.2, max_output_tokens=900)
    return (final or "").strip()

def semantic_select_pages(extracted_text: str, question: str, filename_hint: str = "", top_k: int = 6):
    pages = split_text_into_pages(extracted_text)

    doc_id = f"{filename_hint}|len={len(extracted_text or '')}"
    cache = load_or_build_page_embeddings(doc_id, pages)

    q_emb = _embed_texts([question[:1000]])[0]

    # simple keyword tokens
    q_tokens = set(re.findall(r"\w+", question.lower()))

    scored = []

    for i, emb in enumerate(cache["embeddings"]):
        semantic_score = _cosine(q_emb, emb)

        page_text = cache["pages"][i]["text"].lower()
        page_tokens = set(re.findall(r"\w+", page_text[:3000]))

        # keyword overlap score
        overlap = len(q_tokens & page_tokens)
        keyword_score = overlap / max(len(q_tokens), 1)

        # title page boost (first 3 pages get slight bump)
        page_no = cache["pages"][i]["page"]
        title_boost = 0.1 if page_no <= 3 else 0.0

        final_score = (
            semantic_score * 0.75 +
            keyword_score * 0.20 +
            title_boost * 0.05
        )

        scored.append((final_score, i))

    scored.sort(reverse=True, key=lambda x: x[0])

    picked = []
    for score, idx in scored[:top_k]:
        p = cache["pages"][idx]
        picked.append({
            "page": p["page"],
            "text": p["text"],
            "score": score
        })

    return picked

def _min_default_doc():
    """
    Fallback structure if the model returns incomplete data.
    """
    return {
        "title": "Technical Methodology & Preliminary Project Plan",
        "context_objectives": "This section will be refined upon receipt of the complete RFP details.",
        "methodology": [{
            "phase": "Inception (ECSA Stage 1)",
            "steps": ["Project kickoff", "Information requests", "Stakeholder alignment"]
        }],
        "plan_table": [],
        "project_management": {
            "governance": "Steering Committee monthly.",
            "controls": "Schedule, cost, risk reporting; change control; QA checks."
        },
        "risk": {
            "framework": "Risk management aligned to ISO 31000:2018, using bow-tie thinking where appropriate.",
            "top_risks": [{"name": "Access delays", "treatment": "Early stakeholder engagement"}],
        },
        "quality": {
            "qa_plan": "Quality management aligned to ISO 9001:2015.",
            "controls": "Document control, peer review, stage-gate approvals."
        },

        # NEW SECTIONS
        "deliverables_register": [],
        "document_control": {
            "document_metadata": "Document title, client, project, version, date, author.",
            "review_approval": "Prepared by / Reviewed by / Approved by with sign-off workflow."
        },
        "stakeholder_engagement": {
            "approach": "Stakeholder mapping and engagement aligned with project phases and community sensitivities.",
            "channels": ["Workshops", "Site meetings", "Community liaison", "Formal client progress reporting"]
        },
        "procurement_logic": {
            "stage_gates": [
                "Gate 1: Inception",
                "Gate 2: Concept",
                "Gate 3: Design development",
                "Gate 4: Tender / Construction support"
            ],
            "fee_repricing": "Fee adjustments at agreed stage-gates where scope and quantities are confirmed and agreed."
        },
        "value_engineering": [
            "Optimise design alternatives based on lifecycle cost and constructability.",
            "Rationalise materials and details while maintaining performance.",
        ],

        "additional_services": [],
        "assumptions": ["Site access granted", "Client provides timely data and approvals"],
        "references": ["ECSA Guidelines", "ISO 31000:2018", "ISO 9001:2015"]
    }


def docx_to_markdown(docx_path, max_chars=16000):
    try:
        with open(docx_path, "rb") as f:
            md = mammoth.convert_to_markdown(f).value
        return md[:max_chars]
    except Exception:
        try:
            d = DocxDocument(docx_path)
            text = "\n".join(p.text for p in d.paragraphs)
            return text[:max_chars]
        except Exception:
            return ""


def build_style_and_fewshot(example_docx_path=EXAMPLE_OUTPUT_PATH,
                            example_tender_path=EXAMPLE_TENDER_PATH,
                            tender_chars=6000, docx_chars=12000):
    style_guide = docx_to_markdown(example_docx_path, max_chars=docx_chars)
    try:
        tender_src_text = extract_text_from_pdf(example_tender_path)
    except Exception:
        tender_src_text = ""
    tender_snippet = (tender_src_text or "")[:tender_chars]
    output_snippet = style_guide[:docx_chars]

    return {
        "style_guide": style_guide,
        "example_mapping": {
            "tender_snippet": tender_snippet,
            "output_snippet": output_snippet
        }
    }


def _iter_block_items(parent_doc):
    """Yield each paragraph and table in document order."""
    body = parent_doc._element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent_doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent_doc)


def _is_toc_paragraph(paragraph: Paragraph) -> bool:
    """Detect TOC by field code."""
    p = paragraph._element
    nodes = p.xpath('.//w:fldSimple | .//w:instrText')
    for n in nodes:
        instr = n.get(qn('w:instr')) if n.tag == qn('w:fldSimple') else (n.text or '')
        if instr and "TOC" in instr.upper():
            return True
    txt = (paragraph.text or "").strip().lower()
    return "table of contents" in txt


def _insert_dynamic_toc(doc, STY):
    """
    Insert a fresh, Word-driven Table of Contents.
    Word will populate it when the user updates fields in Word.
    """
    p = doc.add_paragraph()
    p.style = STY["heading1"]
    p.add_run("Table of Contents")

    toc_para = doc.add_paragraph()
    r = toc_para.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')

    instr = OxmlElement('w:instrText')
    instr.text = r'TOC \o "1-3" \h \z \u'

    fld_separate = OxmlElement('w:fldChar')
    fld_separate.set(qn('w:fldCharType'), 'separate')

    placeholder_run = OxmlElement('w:r')
    placeholder_text = OxmlElement('w:t')
    placeholder_text.text = "Right-click here and choose 'Update Field' → 'Update entire table' to build the TOC."
    placeholder_run.append(placeholder_text)

    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')

    r._r.append(fld_begin)
    r._r.append(instr)
    r._r.append(fld_separate)
    r._r.append(placeholder_run)
    r._r.append(fld_end)

    br_para = doc.add_paragraph()
    br_run = br_para.add_run()
    br_run.add_break(WD_BREAK.PAGE)

# ---------------------- DOCX RENDER HELPERS ------------------------
def _find_toc_index(doc: Document):
    for i, p in enumerate(doc.paragraphs):
        txt = (p.text or "").strip()
        if "Table of Contents" in txt or "TABLE OF CONTENTS" in txt:
            return i
    return None


def _delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def _resolve_styles(doc):
    """
    Try preferred style names in order; fall back if missing.
    """
    candidates = {
        "heading1": ["iX Heading 1", "Heading 1"],
        "heading2": ["iX Heading 2", "Heading 2"],
        "heading3": ["iX Heading 3", "Heading 3"],
        "normal":   ["iX Normal", "Normal"],
        "bullet":   ["iX Bullet 1", "Bullet", "List Bullet"],
    }

    found = {}

    def pick(name_list, fallback="Normal"):
        for nm in name_list:
            try:
                _ = doc.styles[nm]
                return nm
            except KeyError:
                continue
        return fallback

    for key, names in candidates.items():
        found[key] = pick(names, "Normal")

    print("[DOCX] Style map ->", found)
    return found


def render_into_template_after_toc(doc_json, template_path, output_path):
    """
    Behaviour:
    - Keep cover page + sign-off / revision page + existing TOC.
    - Find the first 'Introduction' heading.
    - Delete EVERYTHING from that heading onward.
    - Then insert our generated methodology content.
    """
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Branding template not found: {template_path}")

    doc = Document(template_path)
    STY = _resolve_styles(doc)

    # 1) Find "Introduction" cut point
    blocks = list(_iter_block_items(doc))
    cut_idx = None
    intro_pattern = re.compile(r"^\s*(\d+\.?\s*)?introduction\b", re.IGNORECASE)

    for i, blk in enumerate(blocks):
        if isinstance(blk, Paragraph):
            txt = (blk.text or "").strip()
            if not txt:
                continue
            if intro_pattern.match(txt):
                cut_idx = i
                break

    if cut_idx is None:
        print("⚠️ No 'Introduction' heading found; appending methodology at end.")
        cut_idx = len(blocks)
    else:
        print(f"[DOCX] Found 'Introduction' heading at block index {cut_idx}, trimming from there.")

    # 2) Delete from cut_idx onward
    while True:
        blocks = list(_iter_block_items(doc))
        if cut_idx >= len(blocks):
            break
        victim = blocks[cut_idx]
        elm = victim._element
        elm.getparent().remove(elm)

    # 3) Helpers to write content
    def add_heading(text, level=1):
        if not text:
            return
        style_name = STY["heading1"] if level == 1 else STY["heading2"] if level == 2 else STY["heading3"]
        p = doc.add_paragraph()
        p.style = style_name
        p.add_run(str(text))

    def add_text_block(text):
        if text is None:
            return
        if isinstance(text, list):
            text = "\n".join(str(x) for x in text)
        if not isinstance(text, str):
            text = str(text)
        text = text.strip()
        if not text:
            return
        for line in text.split("\n"):
            line = line.strip()
            if line:
                p = doc.add_paragraph()
                p.style = STY["normal"]
                p.add_run(line)

    def add_bullets(items):
        if items is None:
            return
        if isinstance(items, str):
            candidates = [s.strip() for s in re.split(r"[\n;]+", items) if s.strip()]
        elif isinstance(items, dict):
            candidates = [f"{k}: {v}" for k, v in items.items() if v]
        else:
            try:
                candidates = list(items)
            except TypeError:
                candidates = [str(items)]
        for it in candidates:
            it = str(it).strip()
            if not it:
                continue
            p = doc.add_paragraph()
            p.style = STY["bullet"]
            p.add_run(it)

    def _ensure_dict(val, keys):
        if isinstance(val, dict):
            return val
        if isinstance(val, list) and val and isinstance(val[0], dict):
            return val[0]
        d = {k: "" for k in keys}
        if keys:
            d[keys[0]] = val
        return d

    # 4) Insert sections

    # 1. Context & Objectives
    add_heading("1. Context & Objectives (RFP Alignment)", level=1)
    add_text_block(doc_json.get("context_objectives", ""))

    # 2. Detailed Methodology (ECSA Stages)
    add_heading("2. Detailed Methodology (ECSA Stages)", level=1)
    for phase in doc_json.get("methodology", []) or []:
        if isinstance(phase, dict):
            phase_title = phase.get("phase", "")
            steps = phase.get("steps", [])
        else:
            phase_title = str(phase)
            steps = []
        add_heading(phase_title, level=2)
        add_bullets(steps)

    # 3. Preliminary Project Plan (Milestones & Durations)
    add_heading("3. Preliminary Project Plan (Milestones & Durations)", level=1)
    plan = doc_json.get("plan_table", []) or []
    if plan:
        table = doc.add_table(rows=1, cols=8)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr = table.rows[0].cells
        headers = [
            "Work Package",
            "Owner",
            "Start",
            "Finish",
            "Duration (weeks)",
            "Dependencies",
            "Deliverables",
            "Acceptance Criteria",
        ]
        for idx, text in enumerate(headers):
            cell = hdr[idx]
            cell.text = text
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for r in plan:
            if isinstance(r, dict):
                values = [
                    r.get("work_package", ""),
                    r.get("owner", ""),
                    r.get("start", ""),
                    r.get("finish", ""),
                    str(r.get("duration_weeks", "")),
                    r.get("dependencies", ""),
                    r.get("deliverables", ""),
                    r.get("acceptance_criteria", ""),
                ]
            else:
                values = [str(r)] + [""] * 7

            row_cells = table.add_row().cells
            for idx, val in enumerate(values):
                cell = row_cells[idx]
                cell.text = str(val)
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 4. Project Management
    add_heading("4. Project Management: Approach & Methodology", level=1)
    pm = _ensure_dict(doc_json.get("project_management") or {}, ["governance", "controls"])
    add_heading("Governance & Roles", level=2)
    add_text_block(pm.get("governance", ""))
    add_heading("Planning, Controls & Reporting", level=2)
    add_text_block(pm.get("controls", ""))

    # 5. Risk Management
    add_heading("5. Risk Management (ISO 31000:2018)", level=1)
    risk = _ensure_dict(doc_json.get("risk") or {}, ["framework", "top_risks"])
    add_text_block(risk.get("framework", ""))
    top_risks = risk.get("top_risks") or []
    if isinstance(top_risks, dict):
        top_risks = [top_risks]
    elif not isinstance(top_risks, list):
        top_risks = [top_risks]
    risk_bullets = []
    for x in top_risks:
        if isinstance(x, dict):
            risk_bullets.append(f"{x.get('name','')}: {x.get('treatment','')}")
        else:
            risk_bullets.append(str(x))
    add_bullets(risk_bullets)

    # 6. Quality Assurance & Control
    add_heading("6. Quality Assurance & Control (ISO 9001:2015)", level=1)
    qa = _ensure_dict(doc_json.get("quality") or {}, ["qa_plan", "controls"])
    add_text_block(qa.get("qa_plan", ""))
    add_text_block(qa.get("controls", ""))

    # 7. Additional Services
    add_heading("7. Additional Services (ECSA)", level=1)
    add_bullets(doc_json.get("additional_services", []) or [])

    # 8. Deliverables Register
    add_heading("8. Deliverables Register (Mapped to Tender)", level=1)
    reg = doc_json.get("deliverables_register") or []
    if reg:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        headers = ["Deliverable", "ECSA Stage", "Owner", "Format", "Acceptance Criteria"]
        for i, h in enumerate(headers):
            hdr[i].text = h
            for p in hdr[i].paragraphs:
                for run in p.runs:
                    run.bold = True
        for r in reg:
            row = table.add_row().cells
            if isinstance(r, dict):
                row[0].text = str(r.get("deliverable", ""))
                row[1].text = str(r.get("ecsa_stage", ""))
                row[2].text = str(r.get("owner", ""))
                row[3].text = str(r.get("format", ""))
                row[4].text = str(r.get("acceptance_criteria", ""))
            else:
                row[0].text = str(r)
                row[1].text = ""
                row[2].text = ""
                row[3].text = ""
                row[4].text = ""
    else:
        add_text_block("Deliverables will be confirmed upon receipt of the full tender returnables list.")

    # 9. Stakeholder Engagement Strategy
    add_heading("9. Stakeholder Engagement Strategy", level=1)
    se = doc_json.get("stakeholder_engagement") or {}
    if not isinstance(se, dict):
        se = {"approach": str(se), "channels": []}
    add_text_block(se.get("approach", ""))
    add_bullets(se.get("channels", []) or [])

    # 10. Procurement Logic & Stage-Gate Controls
    add_heading("10. Procurement Logic & Stage-Gate Controls", level=1)
    pl = doc_json.get("procurement_logic") or {}
    if not isinstance(pl, dict):
        pl = {"stage_gates": [], "fee_repricing": str(pl)}
    add_bullets(pl.get("stage_gates", []) or [])
    add_text_block(pl.get("fee_repricing", ""))

    # 11. Value Engineering Opportunities
    add_heading("11. Value Engineering Opportunities", level=1)
    add_bullets(doc_json.get("value_engineering", []) or [])

    # 12. Document Control & Approvals
    add_heading("12. Document Control & Approvals", level=1)
    dc = doc_json.get("document_control") or {}
    if not isinstance(dc, dict):
        dc = {"document_metadata": str(dc), "review_approval": ""}
    add_heading("Document Metadata", level=2)
    add_text_block(dc.get("document_metadata", ""))
    add_heading("Review & Approval Workflow", level=2)
    add_text_block(dc.get("review_approval", ""))

    # 13. Assumptions & Dependencies
    add_heading("13. Assumptions & Dependencies", level=1)
    add_bullets(doc_json.get("assumptions", []) or [])

    # 14. Referenced Documents
    add_heading("14. Referenced Documents", level=1)
    add_bullets(doc_json.get("references", []) or [])

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

# --- token clipping ---
def _clip_text_to_tokens(text: str, model: str, max_tokens: int) -> str:
    try:
        import tiktoken
        enc = tiktoken.encoding_for_model(model)
    except Exception:
        import tiktoken
        enc = tiktoken.get_encoding("cl100k_base")
    toks = enc.encode(text or "")
    if len(toks) <= max_tokens:
        return text
    return enc.decode(toks[:max_tokens])


def _strip_code_fences(s: str) -> str:
    if not isinstance(s, str):
        return ""
    return _CODEFENCE_RE.sub("", s).lstrip("\ufeff").strip()


def _guess_first_json(s: str, prefer_array: bool = False) -> str | None:
    if not isinstance(s, str) or not s.strip():
        return None
    s = _strip_code_fences(s)

    try:
        json.loads(s)
        return s
    except Exception:
        pass

    openers = ["[", "{"] if prefer_array else ["{", "["]
    for opener in openers:
        closer = "]" if opener == "[" else "}"
        depth = 0
        start = None
        in_str = False
        esc = False
        for i, ch in enumerate(s):
            if start is None:
                if ch == opener:
                    start, depth = i, 1
                continue
            c = s[i]
            if in_str:
                if esc:
                    esc = False
                elif c == "\\":
                    esc = True
                elif c == '"':
                    in_str = False
            else:
                if c == '"':
                    in_str = True
                elif c == opener:
                    depth += 1
                elif c == closer:
                    depth -= 1
                    if depth == 0:
                        blob = s[start:i + 1]
                        try:
                            json.loads(blob)
                            return blob
                        except Exception:
                            return blob
    return None


def parse_json_loose(s: str | dict | list, key: str | None = None, prefer_array: bool = False):
    if isinstance(s, (dict, list)):
        val = s
    else:
        s = "" if s is None else s
        blob = _guess_first_json(s, prefer_array=prefer_array) or _strip_code_fences(s)
        try:
            val = json.loads(blob)
        except Exception:
            try:
                inner = json.loads(blob)
                if isinstance(inner, str):
                    return parse_json_loose(inner, key=key, prefer_array=prefer_array)
                val = inner
            except Exception:
                try:
                    val = ast.literal_eval(blob)
                except Exception:
                    return {key: None} if key else None

    if key:
        if isinstance(val, dict) and key in val:
            return val
        return {key: val}
    return val


def self_review_json(model, draft_json_str):
    review_prompt = (
        "Review the following JSON for: (1) adherence to the schema, (2) specificity to THIS RFP, "
        "(3) ECSA/ISO anchors present, (4) no regurgitation, (5) clear milestones with durations.\n"
        "If improvements are needed, return an improved JSON. Otherwise, return the original JSON.\n\n"
        f"{draft_json_str}"
    )
    revised_text, _ = call_openai_text(model, review_prompt, temperature=0.2)
    return revised_text or draft_json_str


def _responses_json_block(name, text):
    return (
        f"Return ONLY valid JSON with exactly one top-level key \"{name}\". "
        f"No prose before or after, no code fences.\n"
        f"Schema:\n{{ \"{name}\": <VALUE_FOR_{name.upper()}> }}\n"
        f"No trailing commas.\n\n"
        f"Your task: produce {name}.\n---\n"
    ) + text


def _safe_json_get(blob, key, fallback):
    try:
        return blob.get(key, fallback)
    except Exception:
        return fallback


def _parse_single_key_json(s, key, prefer_array=False):
    return parse_json_loose(s, key=key, prefer_array=prefer_array)

# OCR helpers
def _ocr_status_for_filename(filename: str):
    """Return (status, payload_dict or None) for the OCR job linked to filename."""
    if not filename:
        return None, None
    job_id = FILE_OCR_JOB.get(filename)
    if not job_id:
        return None, None
    with OCR_JOBS_LOCK:
        job = OCR_JOBS.get(job_id)
    if not job:
        return None, None
    status = job.get("status")
    payload = {
        "job_id": job_id,
        "status": status,
        "message": job.get("message", ""),
        "done": job.get("done", 0),
        "total": job.get("total", 0),
        "page": job.get("page", 0),
        "chars": job.get("chars", 0),
        "last_conf": job.get("last_conf", None),
        "progress_url": url_for("ocr_progress", job_id=job_id),
    }
    return status, payload

def _clean_ocr_text(s: str) -> str:
    if not s:
        return ""
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()


def _estimate_skew_angle_pil(gray_img: Image.Image) -> float:
    img = gray_img.copy()
    img.thumbnail((1600, 1600))
    arr = np.array(img)
    thr = np.percentile(arr, 30)
    ys, xs = np.where(arr < thr)

    if len(xs) < 2000:
        return 0.0

    coords = np.column_stack((xs, ys)).astype(np.float32)
    coords -= coords.mean(axis=0)

    cov = np.cov(coords, rowvar=False)
    eigvals, eigvecs = np.linalg.eig(cov)
    principal = eigvecs[:, np.argmax(eigvals)]

    angle = np.degrees(np.arctan2(principal[1], principal[0]))
    if angle < -45:
        angle += 90
    elif angle > 45:
        angle -= 90
    return float(angle)


def preprocess_for_ocr(img: Image.Image) -> Image.Image:
    gray = img.convert("L")
    gray = ImageOps.autocontrast(gray, cutoff=2)
    angle = _estimate_skew_angle_pil(gray)
    if abs(angle) >= 0.4:
        gray = gray.rotate(-angle, expand=True, fillcolor=255)
    gray = gray.filter(ImageFilter.MedianFilter(size=3))
    gray = gray.filter(ImageFilter.UnsharpMask(radius=2, percent=160, threshold=3))

    arr = np.array(gray)
    t = np.percentile(arr, 55)
    bw = (arr > t).astype(np.uint8) * 255
    bw_img = Image.fromarray(bw, mode="L")
    bw_img = bw_img.filter(ImageFilter.MaxFilter(size=3))
    return bw_img

def preprocess_for_ocr_variant(img: Image.Image, variant: int) -> Image.Image:
    """
    Multiple preprocess variants to improve bad pages/tables.
    """
    gray = img.convert("L")
    gray = ImageOps.autocontrast(gray, cutoff=2)

    if variant == 0:
        # current-ish (balanced)
        pass

    elif variant == 1:
        # less aggressive thresholding (helps light text)
        gray = gray.filter(ImageFilter.MedianFilter(size=3))
        gray = gray.filter(ImageFilter.UnsharpMask(radius=2, percent=180, threshold=2))
        arr = np.array(gray)
        t = np.percentile(arr, 62)  # lighter threshold
        bw = (arr > t).astype(np.uint8) * 255
        gray = Image.fromarray(bw, mode="L")

    elif variant == 2:
        # upscale + sharpen (helps small fonts)
        w, h = gray.size
        gray = gray.resize((int(w * 1.5), int(h * 1.5)))
        gray = gray.filter(ImageFilter.UnsharpMask(radius=2, percent=200, threshold=1))
        gray = ImageOps.autocontrast(gray, cutoff=1)

    elif variant == 3:
        # no binarize, just clean grayscale (sometimes best for tables)
        gray = gray.filter(ImageFilter.MedianFilter(size=3))
        gray = ImageOps.autocontrast(gray, cutoff=1)

    return gray

def ocr_page_best_effort(pil_img: Image.Image) -> tuple[str, float]:

    configs = [
        "-l eng --oem 3 --psm 6 -c preserve_interword_spaces=1",
        "-l eng --oem 3 --psm 4 -c preserve_interword_spaces=1",
        "-l eng --oem 3 --psm 11 -c preserve_interword_spaces=1",
        "-l eng --oem 3 --psm 12 -c preserve_interword_spaces=1",  # sparse w/ OSD-ish behavior
        "-l eng --oem 3 --psm 3 -c preserve_interword_spaces=1",
    ]

    best_text = ""
    best_conf = -1.0

    for cfg in configs:
        data = pytesseract.image_to_data(pil_img, config=cfg, output_type=pytesseract.Output.DICT)

        confs = []
        words = []
        n = len(data.get("text", []))
        for i in range(n):
            w = (data["text"][i] or "").strip()
            c = data["conf"][i]
            try:
                c = float(c)
            except Exception:
                c = -1.0
            if w:
                words.append(w)
            if c >= 0:
                confs.append(c)

        mean_conf = float(np.mean(confs)) if confs else 0.0
        text = _clean_ocr_text(" ".join(words))

        if not text:
            raw = pytesseract.image_to_string(pil_img, config=cfg)
            text = _clean_ocr_text(raw)

        score = mean_conf + min(len(text) / 3000.0, 1.0) * 10.0

        if score > best_conf:
            best_conf = score
            best_text = text

    return best_text, best_conf

# ---------- RFP SUMMARY (gpt-5-pro) ----------
def build_rfp_summary(full_text: str, extra_context: str, model: str = "gpt-5-pro") -> dict:
    clipped = _clip_text_to_tokens(full_text or "", model, 85_000)
    extra = f"\nAdditional firm preferences / notes:\n{extra_context}\n" if extra_context else ""

    prompt = _responses_json_block(
        "rfp_summary",
        (
            "You are a senior bid manager and project engineer.\n"
            "Read the FULL RFP text below and produce a compact JSON object named \"rfp_summary\".\n"
            "Focus only on information needed to write a technical methodology and preliminary project plan.\n"
            "Do NOT copy long clauses; summarise clearly.\n\n"
            "JSON structure:\n"
            "{\n"
            "  \"project_context\": string,\n"
            "  \"objectives\": [string],\n"
            "  \"scope_elements\": [string],\n"
            "  \"sites\": [string],\n"
            "  \"beneficiaries\": [string],\n"
            "  \"key_constraints\": [string],\n"
            "  \"deliverables\": [string],\n"
            "  \"stakeholders\": [string],\n"
            "  \"ecsa_stage_hints\": [string],\n"
            "  \"risk_drivers\": [string]\n"
            "}\n\n"
            "Use short bullet-like strings inside arrays.\n"
            f"{extra}\n"
            "=== FULL RFP TEXT ===\n"
            f"{clipped}\n"
        )
    )

    try:
        text, _ = call_openai_text(
            model,
            prompt,
            temperature=0.1,
            max_output_tokens=4500,
            timeout_read=480,
        )
    except Exception as e:
        print(f"[ERROR] Failed to build RFP summary: {e}")
        return {}

    blob = parse_json_loose(text, key="rfp_summary", prefer_array=False)
    summary = blob.get("rfp_summary") if isinstance(blob, dict) else None
    if not isinstance(summary, dict):
        print(f"[WARN] Could not parse rfp_summary; raw head: {str(text)[:400]}")
        return {}
    return summary


# ---------- SECTION-BY-SECTION USING SUMMARY ----------
def generate_methodology_json_by_sections(model: str, full_text: str, extra_context: str):
    """
    1) One call to gpt-5-pro to build a compact rfp_summary.
    2) Multiple small calls (per section) to gpt-5-pro that only see:
       - the rfp_summary JSON
       - a short raw RFP extract (for nuance)
    """
    rfp_summary = build_rfp_summary(full_text, extra_context, model=model)
    summary_text = json.dumps(rfp_summary or {}, ensure_ascii=False, indent=2)

    summary_block = f"\n=== STRUCTURED RFP SUMMARY (rfp_summary) ===\n{summary_text}\n"
    raw_extract = _clip_text_to_tokens(full_text or "", model, 8_000)
    raw_block = f"\n=== RAW RFP EXTRACT (for nuance only) ===\n{raw_extract}\n"

    common_preamble = (
        "ROLE: You are an expert Senior Engineer and Tender Writer.\n"
        "TASK: Produce a comprehensive, high-quality technical methodology for this Tender/RFP.\n"
        "You MUST be tender-specific and use engineering judgment (do NOT regurgitate the scope).\n\n"
        "TONE: Professional, concise, engineering-focused, aligned to ECSA and adjudication scoring.\n\n"
        "REQUIREMENTS:\n"
        "- Expand the technical narrative for each relevant ECSA stage.\n"
        "- Cover sequencing, QA, SHERQ, and community impact.\n"
        "- Include stakeholder engagement strategy.\n"
        "- Include procurement logic (stage-gates, fee re-pricing where applicable).\n"
        "- Incorporate Value Engineering opportunities.\n"
        "- Add compliance statements and referenced standards.\n"
        "- Include document control and approvals.\n"
        "- Risk Management must align to ISO 31000:2018 AND include bow-tie style analysis (textual is fine).\n"
        "- Quality must align to ISO 9001:2015.\n"
    )

    if extra_context:
        common_preamble += f"\nAdditional preferences from the firm:\n{extra_context}\n"

    prompts = {
        "title": (
            common_preamble +
            "Write a short, client-appropriate TITLE for this deliverable (max ~16 words).\n"
            "It should sound like a formal methodology / project plan document.\n"
            + summary_block + raw_block +
            "\n" + _responses_json_block("title", "title: string")
        ),
        "context_objectives": (
            common_preamble +
            "Summarise the context & objectives in 1–2 compact paragraphs.\n"
            "Focus on: site/beneficiaries, problem/opportunity, drivers, success criteria.\n"
            + summary_block + raw_block +
            "\n" + _responses_json_block("context_objectives", "context_objectives: string")
        ),
        "methodology": (
            common_preamble +
            "Provide a Detailed Technical Methodology aligned to ECSA stages.\n"
            "Output as an ARRAY of phases; each phase has a name and 4–10 clear, concise bullet steps.\n"
            + summary_block + raw_block +
            "\n" + _responses_json_block(
                "methodology",
                'methodology: [{"phase": string, "steps": [string]}]'
            )
        ),
        "plan_table": (
            common_preamble +
            "Produce a preliminary project plan with milestones & durations as an ARRAY of rows.\n"
            "Use realistic durations and dependencies; reflect ECSA stages and key deliverables.\n"
            + summary_block + raw_block +
            "\n" + _responses_json_block(
                "plan_table",
                'plan_table: [{"work_package": string, "owner": string, "start": string, "finish": string, '
                '"duration_weeks": number, "dependencies": string, "deliverables": string, "acceptance_criteria": string}]'
            )
        ),
        "project_management": (
            common_preamble +
            "Describe the Project Management approach: governance/roles and controls/reporting.\n"
            "Keep it concise but concrete; align with how the client likely expects to be engaged.\n"
            + summary_block + raw_block +
            "\n" + _responses_json_block(
                "project_management",
                'project_management: {"governance": string, "controls": string}'
            )
        ),
        "risk": (
            common_preamble +
            "Describe Risk Management per ISO 31000:2018.\n"
            "Provide a short framework paragraph and 5–8 top risks with treatments, tailored to this RFP.\n"
            "You may describe bow-tie analysis textually.\n"
            + summary_block + raw_block +
            "\n" + _responses_json_block(
                "risk",
                'risk: {"framework": string, "top_risks": [{"name": string, "treatment": string}]}'
            )
        ),
        "quality": (
            common_preamble +
            "Describe Quality Assurance & Control per ISO 9001:2015.\n"
            "Summarise the QA plan and key controls (reviews, approvals, audits, document control).\n"
            + summary_block + raw_block +
            "\n" + _responses_json_block(
                "quality",
                'quality: {"qa_plan": string, "controls": string}'
            )
        ),
        "additional_services": (
            common_preamble +
            "List additional ECSA-aligned optional services that could be offered beyond the base scope.\n"
            "Only include items that are plausibly relevant to this RFP.\n"
            + summary_block + raw_block +
            "\n" + _responses_json_block("additional_services", 'additional_services: [string]')
        ),
        "assumptions": (
            common_preamble +
            "List explicit assumptions & dependencies underpinning the methodology and plan.\n"
            "These should cover: information, access, approvals, third-party actions, etc.\n"
            + summary_block + raw_block +
            "\n" + _responses_json_block("assumptions", 'assumptions: [string]')
        ),
        "references": (
            common_preamble +
            "List referenced standards/guidelines/documents (bullets).\n"
            "Include ECSA guidelines and relevant ISO standards, plus any obvious national frameworks.\n"
            + summary_block + raw_block +
            "\n" + _responses_json_block("references", 'references: [string]')
        ),

        # NEW SECTIONS
        "deliverables_register": (
            common_preamble +
            "Create a Deliverables Register mapped to this tender.\n"
            "Return an ARRAY of deliverable rows.\n"
            + summary_block + raw_block +
            "\n" + _responses_json_block(
                "deliverables_register",
                'deliverables_register: [{"deliverable": string, "ecsa_stage": string, "owner": string, "format": string, "acceptance_criteria": string}]'
            )
        ),
        "document_control": (
            common_preamble +
            "Provide Document Control and Approval content suitable for a bid submission.\n"
            + summary_block + raw_block +
            "\n" + _responses_json_block(
                "document_control",
                'document_control: {"document_metadata": string, "review_approval": string}'
            )
        ),
        "stakeholder_engagement": (
            common_preamble +
            "Provide a stakeholder engagement strategy tailored to this tender.\n"
            + summary_block + raw_block +
            "\n" + _responses_json_block(
                "stakeholder_engagement",
                'stakeholder_engagement: {"approach": string, "channels": [string]}'
            )
        ),
        "procurement_logic": (
            common_preamble +
            "Provide procurement logic tailored to this tender.\n"
            "Include a clear stage-gate approach and fee re-pricing logic where permitted.\n"
            + summary_block + raw_block +
            "\n" + _responses_json_block(
                "procurement_logic",
                'procurement_logic: {"stage_gates": [string], "fee_repricing": string}'
            )
        ),
        "value_engineering": (
            common_preamble +
            "List 5–10 tender-relevant Value Engineering opportunities.\n"
            "They must be specific to this type of work and context.\n"
            + summary_block + raw_block +
            "\n" + _responses_json_block(
                "value_engineering",
                "value_engineering: [string]"
            )
        ),
    }

    caps = {
        "title": 250,
        "context_objectives": 1600,
        "methodology": 5000,
        "plan_table": 4500,
        "project_management": 1800,
        "risk": 2200,
        "quality": 1800,
        "additional_services": 900,
        "assumptions": 900,
        "references": 700,
        "deliverables_register": 2200,
        "document_control": 900,
        "stakeholder_engagement": 1200,
        "procurement_logic": 900,
        "value_engineering": 900,
    }

    array_keys = {
        "methodology",
        "plan_table",
        "additional_services",
        "assumptions",
        "references",
        "deliverables_register",
        "value_engineering",
    }

    results = {}
    for key, prompt in prompts.items():
        text = ""
        for attempt in range(2):
            try:
                text, _ = call_openai_text(
                    model,
                    prompt,
                    temperature=0.0,
                    max_output_tokens=caps.get(key),
                    timeout_read=480,
                )
                if isinstance(text, str) and text.strip():
                    break
            except Exception as e:
                print(f"[WARN] Error for {key} (attempt {attempt + 1}): {e}")
                time.sleep(2 * (attempt + 1))

        blob = parse_json_loose(text, key=key, prefer_array=(key in array_keys))
        val = blob.get(key) if isinstance(blob, dict) else None
        if val is None and isinstance(text, str):
            print(f"[DEBUG] Raw section head for {key}: {text[:400]}")
        results[key] = val

    base = _min_default_doc()

    merged = {
        "title": results.get("title") or base["title"],
        "context_objectives": results.get("context_objectives") or base["context_objectives"],
        "methodology": results.get("methodology") or base["methodology"],
        "plan_table": results.get("plan_table") or base["plan_table"],
        "project_management": results.get("project_management") or base["project_management"],
        "risk": results.get("risk") or base["risk"],
        "quality": results.get("quality") or base["quality"],
        "additional_services": results.get("additional_services") or base["additional_services"],
        "assumptions": results.get("assumptions") or base["assumptions"],
        "references": results.get("references") or base["references"],

        "deliverables_register": results.get("deliverables_register") or base["deliverables_register"],
        "document_control": results.get("document_control") or base["document_control"],
        "stakeholder_engagement": results.get("stakeholder_engagement") or base["stakeholder_engagement"],
        "procurement_logic": results.get("procurement_logic") or base["procurement_logic"],
        "value_engineering": results.get("value_engineering") or base["value_engineering"],
    }

    return merged

# ---- Pictogram (unchanged from your version, only formatting) ----
def create_methodology_pictogram(doc_json: dict, out_pdf_path: str, out_png_page1_path: str | None = None):
    import os
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib.patches import Rectangle, Circle, Polygon, FancyBboxPatch
    from matplotlib.lines import Line2D
    from textwrap import wrap as _wrap
    from matplotlib.backends.backend_pdf import PdfPages

    GREEN = "#7eb338"
    RED   = "#d6202d"
    GREY  = "#6a6b6c"
    BG    = "#ffffff"
    LEFT_BG = "#f6fbf2"
    LEFT_EDGE = GREY

    def wrap_lines(text, width):
        if text is None:
            return []
        s = str(text).strip()
        if not s:
            return []
        lines = []
        for part in s.split("\n"):
            part = part.strip()
            if not part:
                continue
            lines.extend(_wrap(part, width=width, break_long_words=False, break_on_hyphens=False))
        return [ln for ln in lines if ln.strip()]

    def as_list(x):
        if x is None:
            return []
        if isinstance(x, list):
            return x
        return [x]

    def fit_text_in_box(ax, x, y, text, *, width_chars, max_lines, base_fs, min_fs,
                        color="white", weight="bold", ha="center", va="top"):
        if text is None:
            return
        raw = str(text).strip()
        if not raw:
            return

        lines = wrap_lines(raw, width_chars)
        for fs in range(base_fs, min_fs - 1, -1):
            if len(lines) <= max_lines:
                ax.text(x, y, "\n".join(lines), ha=ha, va=va, fontsize=fs,
                        fontweight=weight, color=color)
                return
            lines = wrap_lines(raw, max(10, width_chars - (base_fs - fs)))

        ax.text(x, y, "\n".join(lines[:max_lines]), ha=ha, va=va, fontsize=min_fs,
                fontweight=weight, color=color)

    methodology = doc_json.get("methodology") or []
    if not isinstance(methodology, list):
        methodology = []

    phases = []
    for i, ph in enumerate(methodology):
        if isinstance(ph, dict):
            name = (ph.get("phase") or f"Stage {i+1}").strip()
            steps = ph.get("steps") or []
        else:
            name = str(ph).strip()
            steps = []
        steps = [str(s).strip() for s in as_list(steps) if str(s).strip()]
        title_lines = wrap_lines(name, width=44)[:4]

        bullet_lines = []
        for s in steps:
            s_lines = wrap_lines(s, width=64)
            if not s_lines:
                continue
            bullet_lines.append(f"• {s_lines[0]}")
            for cont in s_lines[1:]:
                bullet_lines.append(f"  {cont}")

        phases.append({
            "title_lines": title_lines if title_lines else [f"Stage {i+1}"],
            "bullet_lines": bullet_lines if bullet_lines else ["• (no steps provided)"],
        })

    def icon_badge(ax, cx, cy, r, edge=BG):
        ax.add_patch(Circle((cx, cy), r, facecolor="white", edgecolor=edge, lw=1.2))

    def icon_clipboard(ax, cx, cy, r, colour):
        w, h = r * 1.2, r * 1.55
        ax.add_patch(FancyBboxPatch((cx - w/2, cy - h/2), w, h,
                                   boxstyle=f"round,pad=0.01,rounding_size={r*0.15}",
                                   facecolor="none", edgecolor=colour, lw=2.2))
        ax.add_patch(FancyBboxPatch((cx - w*0.22, cy + h*0.33), w*0.44, h*0.16,
                                   boxstyle=f"round,pad=0.01,rounding_size={r*0.12}",
                                   facecolor=colour, edgecolor="none"))
        for k in range(3):
            yy = cy + h*0.15 - k*h*0.22
            ax.add_line(Line2D([cx - w*0.32, cx + w*0.32], [yy, yy], color=colour, lw=2.0))

    def icon_warning(ax, cx, cy, r, colour):
        tri = Polygon([[cx, cy + r*0.95],
                       [cx - r*0.95, cy - r*0.80],
                       [cx + r*0.95, cy - r*0.80]],
                      closed=True, facecolor="none", edgecolor=colour, lw=2.3)
        ax.add_patch(tri)
        ax.add_line(Line2D([cx, cx], [cy + r*0.35, cy - r*0.15], color=colour, lw=2.6))
        ax.add_patch(Circle((cx, cy - r*0.40), r*0.10, facecolor=colour, edgecolor="none"))

    def icon_check(ax, cx, cy, r, colour):
        ax.add_patch(Circle((cx, cy), r*0.92, facecolor="none", edgecolor=colour, lw=2.3))
        ax.add_line(Line2D([cx - r*0.45, cx - r*0.10], [cy - r*0.05, cy - r*0.35], color=colour, lw=2.6))
        ax.add_line(Line2D([cx - r*0.10, cx + r*0.52], [cy - r*0.35, cy + r*0.40], color=colour, lw=2.6))

    def icon_leaf(ax, cx, cy, r, colour):
        leaf = Polygon([[cx, cy + r*0.95],
                        [cx - r*0.85, cy + r*0.25],
                        [cx - r*0.75, cy - r*0.60],
                        [cx, cy - r*0.90],
                        [cx + r*0.85, cy - r*0.10],
                        [cx + r*0.70, cy + r*0.70]],
                       closed=True, facecolor="none", edgecolor=colour, lw=2.3)
        ax.add_patch(leaf)
        ax.add_line(Line2D([cx - r*0.10, cx + r*0.50], [cy - r*0.55, cy + r*0.55], color=colour, lw=2.0))

    right_blocks = [
        ("PROJECT\nMANAGEMENT", icon_clipboard, GREEN),
        ("RISK\nMANAGEMENT", icon_warning, RED),
        ("QUALITY\nASSURANCE", icon_check, GREY),
        ("ENVIRONMENTAL\n& SOCIAL", icon_leaf, GREEN),
    ]

    FIG_W, FIG_H, DPI = 8.27, 11.69, 260
    left_x0, left_x1 = 0.05, 0.70
    right_x0, right_x1 = 0.73, 0.95
    top_y, bottom_y = 0.915, 0.06
    gap = 0.030
    title_font = 10.6
    bullet_font = 9.0

    def px_to_axes_y(fig, px):
        return px / (fig.get_figheight() * fig.dpi)

    def measure_phase_height(fig, ax, phase):
        t1 = ax.text(-10, -10, "\n".join(phase["title_lines"]),
                     fontsize=title_font, fontweight="bold", va="top", ha="left", color=GREY)
        t2 = ax.text(-10, -10, "\n".join(phase["bullet_lines"]),
                     fontsize=bullet_font, va="top", ha="left", color=GREY)
        fig.canvas.draw()
        r = fig.canvas.get_renderer()
        h1_px = t1.get_window_extent(renderer=r).height
        h2_px = t2.get_window_extent(renderer=r).height
        t1.remove()
        t2.remove()

        pad_px = 46
        total_h_axes = px_to_axes_y(fig, h1_px + h2_px + pad_px)
        title_h_axes = px_to_axes_y(fig, h1_px)
        return total_h_axes, title_h_axes

    def draw_page(page_phases, pdf: PdfPages, save_png_path=None):
        fig = plt.figure(figsize=(FIG_W, FIG_H), dpi=DPI, facecolor=BG)
        ax = fig.add_axes([0, 0, 1, 1])
        ax.set_axis_off()
        ax.set_xlim(0, 1)
        ax.set_ylim(0, 1)

        ax.add_patch(Rectangle((0.04, 0.935), 0.92, 0.05, facecolor=GREY, lw=0))
        ax.text(0.5, 0.958, "BID TECHNICAL METHODOLOGY",
                ha="center", va="center", fontsize=16, fontweight="bold", color="white")

        rb_gap = 0.022
        rb_h = ((top_y - bottom_y) - rb_gap * (len(right_blocks) - 1)) / len(right_blocks)

        for j, (label, icon_fn, bg_col) in enumerate(right_blocks):
            y1 = top_y - j * (rb_h + rb_gap)
            y0 = y1 - rb_h

            ax.add_patch(FancyBboxPatch(
                (right_x0, y0), right_x1 - right_x0, rb_h,
                boxstyle="round,pad=0.006,rounding_size=0.012",
                facecolor=bg_col, edgecolor=BG, lw=0.0
            ))

            cx = (right_x0 + right_x1) / 2
            icon_y = y1 - rb_h * 0.30
            badge_r = 0.040
            icon_badge(ax, cx, icon_y, badge_r, edge=BG)
            icon_fn(ax, cx, icon_y, r=0.030, colour=bg_col)

            label_y = y0 + rb_h * 0.38
            fit_text_in_box(
                ax, cx, label_y, label,
                width_chars=14, max_lines=3,
                base_fs=10, min_fs=7,
                color="white", weight="bold",
                ha="center", va="top"
            )

        measures = [measure_phase_height(fig, ax, p) for p in page_phases]
        heights = [m[0] for m in measures]
        title_heights = [m[1] for m in measures]

        y_cursor = top_y
        for i, (p, box_h) in enumerate(zip(page_phases, heights)):
            y1 = y_cursor
            y0 = y1 - box_h

            ax.add_patch(FancyBboxPatch(
                (left_x0, y0), left_x1 - left_x0, box_h,
                boxstyle="round,pad=0.006,rounding_size=0.012",
                facecolor=LEFT_BG, edgecolor=LEFT_EDGE, lw=1.1
            ))

            ax.add_patch(Rectangle((left_x0, y0), 0.012, box_h, facecolor=RED, lw=0))

            TOP_PAD = 0.010
            LEFT_PAD = 0.020
            TITLE_BULLET_GAP = 0.010

            inner_top = y1 - TOP_PAD

            tx = left_x0 + 0.030
            ty = inner_top
            ax.text(tx, ty, "\n".join(p["title_lines"]),
                    ha="left", va="top", fontsize=title_font, fontweight="bold", color=GREY)

            title_h = title_heights[i]
            by = ty - title_h - TITLE_BULLET_GAP
            ax.text(left_x0 + LEFT_PAD, by, "\n".join(p["bullet_lines"]),
                    ha="left", va="top", fontsize=bullet_font, color=GREY)

            if i < len(page_phases) - 1:
                ax.text((left_x0 + left_x1) / 2, y0 - gap / 2,
                        "↓", ha="center", va="center", fontsize=18, color=RED)

            y_cursor = y0 - gap

        ax.text(0.5, 0.028, "iX Engineers | Bid Technical Methodology Pictogram",
                ha="center", va="center", fontsize=8.5, color=GREY)

        pdf.savefig(fig, facecolor=BG)
        if save_png_path:
            fig.savefig(save_png_path, facecolor=BG)
        plt.close(fig)

    os.makedirs(os.path.dirname(out_pdf_path), exist_ok=True)

    with PdfPages(out_pdf_path) as pdf:
        page = []
        scratch = plt.figure(figsize=(FIG_W, FIG_H), dpi=DPI, facecolor=BG)
        ax_s = scratch.add_axes([0, 0, 1, 1])
        ax_s.set_axis_off()
        ax_s.set_xlim(0, 1)
        ax_s.set_ylim(0, 1)

        available_h = (top_y - bottom_y)
        used = 0.0

        for p in phases:
            h, _ = measure_phase_height(scratch, ax_s, p)
            needed = h + (gap if page else 0.0)

            if page and (used + needed) > available_h:
                draw_page(
                    page,
                    pdf,
                    save_png_path=out_png_page1_path if out_png_page1_path and pdf.get_pagecount() == 0 else None
                )
                page = []
                used = 0.0

            page.append(p)
            used += needed

        if page:
            draw_page(
                page,
                pdf,
                save_png_path=out_png_page1_path if out_png_page1_path and pdf.get_pagecount() == 0 else None
            )

        plt.close(scratch)

def extract_text_from_image(filepath):
    img = Image.open(filepath).convert("L")
    config = "--oem 3 --psm 6"
    return (pytesseract.image_to_string(img, config=config) or "").strip()


def extract_text_from_file(filepath):
    ext = os.path.splitext(filepath.lower())[1]
    if ext == ".pdf":
        return extract_text_from_pdf(filepath, ocr_if_small=False)
    if ext in [".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"]:
        return extract_text_from_image(filepath)
    return ""


def _looks_like_real_text(s: str) -> bool:
    if not s:
        return False
    s = s.strip()
    if len(s) < 800:
        return False
    alnum = sum(ch.isalnum() for ch in s)
    ratio = alnum / max(1, len(s))
    if ratio < 0.20:
        return False
    spaces = s.count(" ")
    if spaces < 30:
        return False
    return True

# ---------------------- ROUTES ---------------------------------

@app.route('/')
def index():
    return render_template('index.html')


AUTHORIZED_EMAILS = ['anathi.c@ixengineers.co.za', 'tumi.l@ixengineers.co.za', 'kgwerano.c@ixengineers.co.za', 'adrian.c@ixengineers.co.za']


@app.route('/login')
def login():
    auth_url = msal_app.get_authorization_request_url(SCOPES, redirect_uri=redirect_uri)
    return redirect(auth_url)


@app.route('/welcome/getAToken')
def get_token():
    code = request.args.get('code')
    result = msal_app.acquire_token_by_authorization_code(code, SCOPES, redirect_uri=redirect_uri)

    if "access_token" in result:
        session['access_token'] = result['access_token']
        user_profile = get_user_profile(result['access_token'])
        session['user'] = {
            'email': user_profile.get('mail', user_profile.get('userPrincipalName'))
        }
        return redirect(url_for('welcome'))
    else:
        return jsonify(result.get("error_description")), 400


def get_name(email):
    return email.split('.')[0].capitalize()


@app.route('/welcome')
def welcome():
    if 'user' in session:
        user_email = session['user']['email']
        user_name = get_name(user_email)
        return render_template('welcome.html', user_name=user_name, user_email=user_email)
    else:
        return redirect(url_for('login'))


@app.route('/dashboard')
def dashboard():
    return render_template('dashboard3.html')


@app.route('/dashboardTwo')
def dashboardTwo():
    return render_template('dashboard.html')

def get_user_profile(access_token):
    graph_url = "https://graph.microsoft.com/v1.0/me"
    headers = {
        'Authorization': f'Bearer {access_token}'
    }
    response = requests.get(graph_url, headers=headers)
    return response.json()


@app.route('/plan/<filename>')
def plan(filename):
    return render_template('generate_plan.html', pdf_filename=filename)


@app.route('/plan')
def plan_blank():
    filename = request.args.get('filename', '')
    return render_template('generate_plan.html', pdf_filename=filename)


@app.route('/feedback', methods=['GET', 'POST'])
def feedback():
    if request.method == 'POST':
        user_email = session['user']['email']
        name = get_name(user_email)
        feedback_text = request.form['feedback']
        rating_str = request.form['rating']
        timestamp = datetime.now(south_africa_tz).strftime('%Y-%m-%d %H:%M:%S')

        rating_map = {
            '1 - Very Dissatisfied': 1,
            '2 - Dissatisfied': 2,
            '3 - Neutral': 3,
            '4 - Satisfied': 4,
            '5 - Very Satisfied': 5
        }
        rating = rating_map.get(rating_str, None)

        if rating is not None:
            new_feedback = Feedback(
                name=user_email,
                feedback_text=feedback_text,
                rating=rating,
                timestamp=timestamp
            )
            db.session.add(new_feedback)
            db.session.commit()
        else:
            return jsonify({'message': 'Invalid rating'}), 400

        return redirect(url_for('dashboard'))

    user_email = session['user']['email']
    name = get_name(user_email)
    return render_template('feedback.html', name=name, user_email=user_email, AUTHORIZED_EMAILS=AUTHORIZED_EMAILS)


@app.route('/ask/<filename>', methods=['GET'])
def ask(filename):
    return render_template('ask_question.html', pdf_filename=filename)


MAX_AUTO_PROCESS_SIZE = 1.01 * 1024 * 1024  # 10 MB


@app.route("/upload", methods=["POST"])
def upload_files():
    if 'file' not in request.files:
        return jsonify({"message": "No file part in the request."}), 400

    files = request.files.getlist('file')
    if not files:
        return jsonify({"message": "No files uploaded."}), 400

    saved_files = []
    for file in files:
        if file.filename == '':
            continue
        filename = secure_filename(file.filename)
        save_path = os.path.join("uploads", filename)
        file.save(save_path)
        saved_files.append(filename)
        ext = os.path.splitext(filename.lower())[1]
        if ext in [".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"]:
            text = extract_text_from_file(save_path)
            text_filepath = os.path.join("temp_data", filename + ".txt")
            with open(text_filepath, "w", encoding="utf-8") as f:
                f.write(text)

    if not saved_files:
        return jsonify({"message": "No valid files uploaded."}), 400

    return jsonify({
        "message": f"{len(saved_files)} file(s) uploaded successfully.",
        "filenames": saved_files
    })


@app.route('/merge', methods=['POST'])
def merge_pdfs():
    data = request.get_json(silent=True) or {}
    files = data.get('files', [])
    output_name = (data.get('output_name') or '').strip()

    # ✅ FIX: define scanned flag from frontend
    scanned = bool(data.get('is_scanned', False))

    if not files or any(not f or not isinstance(f, str) for f in files):
        return jsonify({"message": "No valid files provided"}), 400

    for f in files:
        path = os.path.join(app.config['UPLOAD_FOLDER'], f)
        if not os.path.isfile(path):
            return jsonify({"message": f"File not found: {f}"}), 400

    # Sanitize user-specified name or fallback
    if output_name:
        if not output_name.lower().endswith(".pdf"):
            output_name += ".pdf"
        output_name = secure_filename(output_name)
    else:
        merged_basename = "_".join([os.path.splitext(f)[0] for f in files])
        output_name = f"{merged_basename}_{uuid.uuid4().hex[:6]}.pdf"

    merged_pdf_path = os.path.join("uploads", output_name)

    # -----------------------
    # SINGLE FILE
    # -----------------------
    if len(files) == 1:
        filename = secure_filename(files[0])
        pdf_path = os.path.join("uploads", filename)
        text_filepath = os.path.join("temp_data", filename + ".txt")

        if not os.path.exists(pdf_path):
            return jsonify({"message": f"File not found: {filename}"}), 404

        # If we already have a decent text file, return immediately
        if os.path.exists(text_filepath) and os.path.getsize(text_filepath) > 2000:
            return jsonify({
                "message": "Single document ready (cached text).",
                "filename": filename,
                "ocr": {"status": "done", "job_id": FILE_OCR_JOB.get(filename)}
            })

        # Not scanned: do embedded extraction only (no OCR)
        if not scanned:
            try:
                out_parts = []
                with fitz.open(pdf_path) as doc:
                    for i in range(len(doc)):
                        page = doc.load_page(i)
                        embedded = (page.get_text("text") or "").strip()
                        if embedded:
                            out_parts.append(f"=== PAGE {i+1} (embedded) ===\n{embedded}\n")

                with open(text_filepath, "w", encoding="utf-8") as f:
                    f.write(_clean_ocr_text("\n\n".join(out_parts)))

            except Exception as e:
                return jsonify({"message": f"Failed to extract embedded text: {e}"}), 500

            FILE_OCR_JOB.pop(filename, None)

            return jsonify({
                "message": "Single document ready (embedded text).",
                "filename": filename,
                "ocr": {"status": "skipped", "job_id": None}
            })

        # Scanned: start OCR background job
        job_id = uuid.uuid4().hex[:12]
        with OCR_JOBS_LOCK:
            OCR_JOBS[job_id] = {
                "status": "queued",
                "message": "Queued for OCR",
                "done": 0,
                "total": 0,
                "page": 0,
                "filename": filename,
                "pdf_path": pdf_path,
                "text_filepath": text_filepath,
                "created_at": datetime.utcnow().isoformat(),
            }
        FILE_OCR_JOB[filename] = job_id

        t = threading.Thread(
            target=_ocr_pdf_background,
            args=(job_id, pdf_path, text_filepath),
            daemon=True
        )
        t.start()

        return jsonify({
            "message": "OCR started.",
            "filename": filename,
            "ocr": {
                "job_id": job_id,
                "status": "queued",
                "progress_url": url_for("ocr_progress", job_id=job_id),
            }
        })

    # -----------------------
    # MULTI FILE MERGE (2+)
    # -----------------------
    if len(files) < 2:
        return jsonify({"message": "At least two files are required to merge."}), 400

    merger = PdfMerger()
    for file in files:
        safe = secure_filename(file)
        path = os.path.join("uploads", safe)
        if not os.path.exists(path):
            return jsonify({"message": f"File not found: {file}"}), 404
        merger.append(path)

    with open(merged_pdf_path, "wb") as f:
        merger.write(f)
    merger.close()

    merged_txt_path = os.path.join("temp_data", output_name + ".txt")

    # If NOT scanned, do normal embedded extraction (your existing flow)
    if not scanned:
        try:
            text = extract_text_from_file(merged_pdf_path)
        except Exception as e:
            return jsonify({"message": f"Failed to extract text: {str(e)}"}), 500

        with open(merged_txt_path, "w", encoding="utf-8") as f:
            f.write(text or "")

        FILE_OCR_JOB.pop(output_name, None)

        print("[MERGE] Merged file:", output_name)
        print("[MERGE] Writing txt:", merged_txt_path, "chars:", len(text or ""))

        return jsonify({
            "message": "Merge successful (embedded text).",
            "filename": output_name,
            "ocr": {"status": "skipped", "job_id": None}
        })

    # If scanned, start OCR on the merged PDF so the UI can wait correctly
    job_id = uuid.uuid4().hex[:12]
    with OCR_JOBS_LOCK:
        OCR_JOBS[job_id] = {
            "status": "queued",
            "message": "Queued for OCR",
            "done": 0,
            "total": 0,
            "page": 0,
            "filename": output_name,
            "pdf_path": merged_pdf_path,
            "text_filepath": merged_txt_path,
            "created_at": datetime.utcnow().isoformat(),
        }
    FILE_OCR_JOB[output_name] = job_id

    t = threading.Thread(
        target=_ocr_pdf_background,
        args=(job_id, merged_pdf_path, merged_txt_path),
        daemon=True
    )
    t.start()

    print("[MERGE] Merged file:", output_name)
    print("[MERGE] OCR started for merged PDF. job:", job_id)

    return jsonify({
        "message": "Merge successful. OCR started.",
        "filename": output_name,
        "ocr": {
            "job_id": job_id,
            "status": "queued",
            "progress_url": url_for("ocr_progress", job_id=job_id),
        }
    })

@app.route('/uploads/<filename>')
def uploaded_file(filename):
    filename = secure_filename(filename)
    path = os.path.join(app.config['UPLOAD_FOLDER'], filename)

    if not os.path.exists(path):
        return jsonify({'message': f'File not found: {filename}'}), 404

    download = request.args.get("download", "0") == "1"

    return send_file(
        path,
        as_attachment=download,
        download_name=filename
    )


@app.route('/_test_render')
def _test_render():
    dummy = _min_default_doc()
    out_name = secure_filename("TEST_output.docx")
    out_path = os.path.join(app.config['UPLOAD_FOLDER'], out_name)
    try:
        render_into_template_after_toc(dummy, COMPANY_TEMPLATE_PATH, out_path)
    except Exception as e:
        return f"Render error: {e}", 500
    return f"OK -> /uploads/{out_name}"


@app.route("/generate_answers", methods=["POST"])
def generate_answers():
    data = request.json
    filename = data.get("filename")

    with open(f"temp_data/{filename}.txt", "r", encoding="utf-8") as f:
        extracted_text = f.read()

    answers = {
        "What is scope of this project?": get_answer_from_openai("What is scope of this project?", extracted_text),
        "Where is the Work location?": get_answer_from_openai("Where is the Works location?", extracted_text),
        "What is the closing date?": get_answer_from_openai("What is the closing date?", extracted_text),
        "What is the closing time?": get_answer_from_openai("What is the closing time?", extracted_text),
        "Where is the tender submission box?": get_answer_from_openai("Where is the tender submission box?", extracted_text),
        "Is the evaluation criteria 80/20 or 90/10?": get_answer_from_openai("Is the evaluation criteria 80/20 or 90/10?", extracted_text),
        "What is the required Key Personnel?": get_answer_from_openai("What is the required Key Personnel?", extracted_text),
        "What is the required Company Experience?": get_answer_from_openai("What is the required Company Experience?", extracted_text),
        "Does the Tender require reference letters?": get_answer_from_openai("Does the Tender require reference letters?", extracted_text),
        "Who will be our competitors?": get_answer_from_openai("Who will be our competitors?", extracted_text),
    }

    return jsonify({
        "answers": answers
    })


@app.route('/ask', methods=['POST'])
def ask_pdf():
    data = request.get_json(silent=True) or {}
    question = (data.get('question') or '').strip()
    filename = (data.get('filename') or '').strip()
    mode = (data.get('mode') or '').strip().lower()   # "general" for dashboard.html

    if not filename:
        return jsonify({'answer': 'Filename not provided.'}), 400
    if not question:
        return jsonify({'answer': 'Question not provided.'}), 400

    # ✅ OCR gating (keep as-is)
    status, ocr_payload = _ocr_status_for_filename(filename)
    if status in ("queued", "running"):
        return jsonify({
            "answer": "OCR still in progress. Please wait for it to finish, then ask again.",
            "ocr": ocr_payload
        }), 202

    txt_path = os.path.join("temp_data", filename + ".txt")
    if not os.path.exists(txt_path):
        return jsonify({'answer': 'PDF text file not found.'}), 404

    with open(txt_path, "r", encoding="utf-8") as f:
        text = f.read() or ""

    # ✅ dashboard.html: use old behavior + whole-doc summary when asked
    if mode == "general":
        answer = answer_general_or_summary(question, text)
        return jsonify({'answer': answer})

    # otherwise keep your existing tender-style / semantic retrieval mode:
    answer = get_answer_from_openai(question, text, filename_hint=filename)
    return jsonify({'answer': answer})


def extract_text_from_pdf(filepath, ocr_if_small=True, ocr_dpi=400, max_pages=None):
    """
    Drop-in replacement.

    Key improvements:
    - Hybrid extraction: ALWAYS keep embedded text if present, and OCR additionally when embedded is weak.
    - Page markers to enable better downstream retrieval (=== PAGE n ===).
    - Slightly safer fallbacks and cleaning.
    - Higher default DPI (400) for scanned docs (you can pass 300 if you want).

    NOTE: This function is used by your merge flow for multi-file merges.
          Your single-file background OCR uses _ocr_pdf_background().
    """
    out_parts = []

    with fitz.open(filepath) as doc:
        page_count = len(doc)

        for i in range(page_count):
            if max_pages is not None and i >= max_pages:
                break

            page = doc.load_page(i)

            # ---- Embedded text (keep it if present, even if short) ----
            embedded_page = (page.get_text("text") or "").strip()
            if embedded_page:
                out_parts.append(f"=== PAGE {i+1} (embedded) ===\n{embedded_page}\n")

            # Decide whether we should OCR in addition
            if not ocr_if_small:
                continue

            # If embedded text looks strong, skip OCR for speed/cost
            if _looks_like_real_text(embedded_page):
                continue

            # ---- OCR path (best-effort) ----
            zoom = float(ocr_dpi) / 72.0
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img = Image.open(io.BytesIO(pix.tobytes("png")))

            # Variant 0: your main preprocess
            pre0 = preprocess_for_ocr(img)
            txt, conf = ocr_page_best_effort(pre0)

            # If weak, try multiple preprocess variants (especially useful for tables)
            if (len(txt) < 300) or (conf < 33.0):
                best_txt, best_conf = txt, conf

                # Try your variants 1..3 (1: lighter threshold, 2: upscale+sharpen, 3: grayscale/no binarize)
                for v in (1, 2, 3):
                    pre_v = preprocess_for_ocr_variant(img, v)
                    t2, c2 = ocr_page_best_effort(pre_v)

                    # score: confidence + content length bonus
                    score_best = best_conf + min(len(best_txt) / 2500.0, 1.0) * 10.0
                    score_new  = c2        + min(len(t2)       / 2500.0, 1.0) * 10.0

                    if score_new > score_best:
                        best_txt, best_conf = t2, c2

                txt, conf = best_txt, best_conf

            txt = (txt or "").strip()
            if txt:
                out_parts.append(f"=== PAGE {i+1} (ocr) ===\n{txt}\n")

            # Helpful logging
            print(f"[OCR] page {i+1}/{page_count} chars={len(txt)} conf={conf:.1f}")

    final = _clean_ocr_text("\n\n".join([p for p in out_parts if p]))
    return final

def run_ocr_job(job_id: str, pdf_path: str, text_filepath: str, filename: str):
    try:
        OCR_JOBS[job_id] = {"status": "running", "done": 0, "total": 0, "chars": 0, "message": "Starting OCR...", "filename": filename}

        with fitz.open(pdf_path) as doc:
            total = len(doc)
            OCR_JOBS[job_id]["total"] = total

            parts = []
            for i in range(total):
                page = doc.load_page(i)
                zoom = 300 / 72.0
                mat = fitz.Matrix(zoom, zoom)
                pix = page.get_pixmap(matrix=mat, alpha=False)
                img = Image.open(io.BytesIO(pix.tobytes("png")))

                pre = preprocess_for_ocr(img)
                txt, conf = ocr_page_best_effort(pre)

                parts.append(txt or "")
                combined = "\n\n".join(parts)

                with open(text_filepath, "w", encoding="utf-8") as f:
                    f.write(_clean_ocr_text(combined))

                OCR_JOBS[job_id].update({
                    "done": i + 1,
                    "chars": len(combined),
                    "message": f"OCR page {i+1}/{total} (conf {conf:.1f})"
                })

        OCR_JOBS[job_id]["status"] = "done"
        OCR_JOBS[job_id]["message"] = "OCR complete."

    except Exception as e:
        OCR_JOBS[job_id]["status"] = "error"
        OCR_JOBS[job_id]["message"] = f"OCR failed: {e}"


def _ocr_pdf_background(job_id: str, pdf_path: str, text_filepath: str, ocr_dpi: int = 300):
    try:
        with OCR_JOBS_LOCK:
            OCR_JOBS[job_id]["status"] = "running"
            OCR_JOBS[job_id]["message"] = "Starting OCR..."

        os.makedirs(os.path.dirname(text_filepath), exist_ok=True)
        with open(text_filepath, "w", encoding="utf-8") as f:
            f.write("")

        with fitz.open(pdf_path) as doc:
            total = len(doc)
            with OCR_JOBS_LOCK:
                OCR_JOBS[job_id]["total"] = total
                OCR_JOBS[job_id]["done"] = 0

            for i in range(total):
                with OCR_JOBS_LOCK:
                    if OCR_JOBS[job_id].get("cancel"):
                        OCR_JOBS[job_id]["status"] = "cancelled"
                        OCR_JOBS[job_id]["message"] = "Cancelled"
                        return

                page = doc.load_page(i)
                zoom = ocr_dpi / 72.0
                mat = fitz.Matrix(zoom, zoom)
                pix = page.get_pixmap(matrix=mat, alpha=False)
                img = Image.open(io.BytesIO(pix.tobytes("png")))

                candidates = []
                for v in (0, 1, 2, 3):
                    pre_v = preprocess_for_ocr_variant(img, v)
                    t, c = ocr_page_best_effort(pre_v)
                    score = c + min(len(t) / 2500.0, 1.0) * 10.0
                    candidates.append((score, t, c, v))

                # pick best
                candidates.sort(key=lambda x: x[0], reverse=True)
                _, txt, conf, best_v = candidates[0]

                with open(text_filepath, "a", encoding="utf-8") as f:
                    f.write(f"\n\n=== PAGE {i+1} ===\n")
                    f.write((txt or "") + "\n")

                with OCR_JOBS_LOCK:
                    OCR_JOBS[job_id]["done"] = i + 1
                    OCR_JOBS[job_id]["page"] = i + 1
                    OCR_JOBS[job_id]["last_conf"] = float(conf)
                    OCR_JOBS[job_id]["message"] = f"OCR page {i+1}/{total}"
                    try:
                        OCR_JOBS[job_id]["chars"] = os.path.getsize(text_filepath)
                    except Exception:
                        pass

                print(f"[OCR] job={job_id} page {i+1}/{total} chars={len(txt or '')} conf={conf:.1f}")

        with OCR_JOBS_LOCK:
            OCR_JOBS[job_id]["status"] = "done"
            OCR_JOBS[job_id]["message"] = "OCR complete"
            OCR_JOBS[job_id]["completed_at"] = datetime.utcnow().isoformat()

    except Exception as e:
        print("[OCR] Background OCR error:", e)
        with OCR_JOBS_LOCK:
            OCR_JOBS[job_id]["status"] = "error"
            OCR_JOBS[job_id]["message"] = str(e)

def find_snippets(text: str, query: str, window: int = 800, max_hits: int = 3):
    if not text or not query:
        return []
    q = re.escape(query.strip())
    hits = []
    for m in re.finditer(q, text, flags=re.IGNORECASE):
        start = max(0, m.start() - window)
        end = min(len(text), m.end() + window)
        hits.append(text[start:end])
        if len(hits) >= max_hits:
            break
    return hits


# ===================== SEMANTIC RETRIEVAL (EMBEDDINGS) =====================
_EMBED_MODEL = os.getenv("OPENAI_EMBED_MODEL", "text-embedding-3-small")
_EMBED_DIR = os.path.join("temp_data", ".embeddings")
os.makedirs(_EMBED_DIR, exist_ok=True)

def _doc_key_from_context(context: str, filename: Optional[str] = None) -> str:
    """Stable-ish key for caching embeddings."""
    if filename:
        base = filename
    else:
        s = (context or "")
        sample = (s[:20000] + "\n...\n" + s[-20000:]) if len(s) > 45000 else s
        base = hashlib.sha1(sample.encode("utf-8", errors="ignore")).hexdigest()
    return re.sub(r"[^a-zA-Z0-9._-]+", "_", base)

def _embedding_api(texts: List[str]) -> List[List[float]]:
    """Call OpenAI embeddings via HTTPS."""
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY env var not set")

    url = "https://api.openai.com/v1/embeddings"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    payload = {"model": _EMBED_MODEL, "input": texts}

    s = _get_openai_session()
    r = s.post(url, headers=headers, json=payload, timeout=(10, 120))
    if r.status_code >= 400:
        raise RuntimeError(f"Embeddings error {r.status_code}: {r.text}")

    data = r.json()
    out = []
    for item in (data.get("data") or []):
        out.append(item.get("embedding") or [])
    if len(out) != len(texts):
        raise RuntimeError("Embeddings API returned unexpected shape.")
    return out

def _chunk_text_for_index(text: str, chunk_size: int = 2600, overlap: int = 300) -> List[str]:
    text = text or ""
    if not text.strip():
        return []
    chunks = []
    step = max(1, chunk_size - overlap)
    i = 0
    while i < len(text):
        chunk = text[i:i + chunk_size]
        if chunk.strip():
            chunks.append(chunk)
        i += step
    return chunks

def _load_or_build_embedding_index(full_text: str, filename: Optional[str] = None,
                                   chunk_size: int = 2600, overlap: int = 300):
    """Returns (chunks, vectors_np). Caches to disk."""
    doc_key = _doc_key_from_context(full_text, filename)
    cache_path = os.path.join(_EMBED_DIR, f"{doc_key}.pkl")

    try:
        if os.path.exists(cache_path):
            with open(cache_path, "rb") as f:
                obj = pickle.load(f)
            chunks = obj.get("chunks") or []
            vecs = obj.get("vecs")
            if isinstance(chunks, list) and isinstance(vecs, np.ndarray) and len(chunks) == vecs.shape[0]:
                return chunks, vecs
    except Exception:
        pass

    chunks = _chunk_text_for_index(full_text, chunk_size=chunk_size, overlap=overlap)
    if not chunks:
        return [], np.zeros((0, 1), dtype=np.float32)

    vec_list = []
    B = 96
    for i in range(0, len(chunks), B):
        batch = chunks[i:i + B]
        embs = _embedding_api(batch)
        vec_list.extend(embs)

    vecs = np.array(vec_list, dtype=np.float32)
    norms = np.linalg.norm(vecs, axis=1, keepdims=True) + 1e-12
    vecs = vecs / norms

    try:
        with open(cache_path, "wb") as f:
            pickle.dump({"chunks": chunks, "vecs": vecs}, f, protocol=pickle.HIGHEST_PROTOCOL)
    except Exception as e:
        print("[EMBED] cache save failed:", e)

    return chunks, vecs

def select_relevant_chunks(full_text: str, question: str,
                           chunk_size: int = 4500, overlap: int = 700, top_k: int = 6,
                           filename: Optional[str] = None):
    """
    Semantic-first chunk selection.

    - Uses OpenAI embeddings to retrieve chunks by meaning (handles synonyms / paraphrases).
    - Falls back to lightweight keyword scoring if embeddings fail.
    """
    text = full_text or ""
    q = (question or "").strip()
    if not text.strip():
        return [""]

    # 1) Semantic retrieval (preferred)
    try:
        idx_chunks, idx_vecs = _load_or_build_embedding_index(
            text,
            filename=filename,
            chunk_size=min(2600, max(1200, chunk_size // 2)),
            overlap=min(300, max(120, overlap // 3)),
        )
        if idx_chunks and idx_vecs.size:
            q_vec = np.array(_embedding_api([q])[0], dtype=np.float32)
            q_vec = q_vec / (np.linalg.norm(q_vec) + 1e-12)

            sims = idx_vecs @ q_vec
            top_idx = np.argsort(-sims)[:max(top_k, 1)]
            picked = [idx_chunks[int(i)] for i in top_idx if float(sims[int(i)]) > 0.05]
            if picked:
                return picked[:top_k]
    except Exception as e:
        print("[EMBED] semantic retrieval failed; falling back to keyword:", e)

    # 2) Keyword fallback
    ql = q.lower()
    tokens = set(re.findall(r"[a-z0-9][a-z0-9\-\/\.]{1,}", ql))
    if not tokens:
        return [text[:chunk_size]]

    chunks = []
    step = max(1, chunk_size - overlap)
    i = 0
    while i < len(text):
        chunks.append(text[i:i + chunk_size])
        i += step

    scored = []
    for c in chunks:
        cl = c.lower()
        score = 0
        for t in tokens:
            if t in cl:
                score += 2 if len(t) >= 5 else 1
        scored.append((score, c))

    scored.sort(key=lambda x: x[0], reverse=True)
    best = [c for s, c in scored[:top_k] if s > 0]
    return best or [text[:chunk_size]]

def get_answer_from_openai(question: str, extracted_text: str, filename_hint: str = "") -> str:
    """
    Semantic retrieval + page references.
    """
    try:
        model = os.getenv("OPENAI_QA_MODEL", "gpt-4o-mini")

        # 1) semantic retrieval (synonyms handled by embeddings)
        picked_pages = semantic_select_pages(extracted_text, question, filename_hint=filename_hint, top_k=7)

        # build context with page markers
        ctx_parts = []
        ref_pages = []
        for p in picked_pages:
            ref_pages.append(p["page"])
            ctx_parts.append(f"=== PAGE {p['page']} ===\n{p['text']}\n")
        combined = "\n\n---\n\n".join(ctx_parts)

        # detect location-type questions (allow "where in doc" style answers)
        q_lower = (question or "").lower()
        location_question = any(
            kw in q_lower for kw in [
                "where", "which page", "what page", "which section", "what section",
                "locate", "find", "reference", "clause"
            ]
        )

        rules = f"""
You answer questions using ONLY the provided CONTEXT.

OUTPUT FORMAT (mandatory):
Answer: <direct answer or 'Not found in extracted text.'>
References: Page(s) <comma-separated page numbers> OR '(none)'
Evidence: <1–2 short exact quotes from CONTEXT> OR '(none)'

RULES:
- Give the DIRECT answer first (date/time/value/name/yes-no).
- If the answer is not explicitly present in CONTEXT, output:
  Answer: Not found in extracted text.
- Do NOT say "refer to/see page..." as the answer unless the user asked WHERE something is located.
  Location-style answer allowed: {str(location_question)}
- Evidence quotes must be exact snippets from CONTEXT.
"""

        prompt = f"""{rules}

CONTEXT:
{combined}

QUESTION: {question}
"""

        text, _ = call_openai_text(model, prompt, temperature=0.0, max_output_tokens=450)
        out = (text or "").strip()

        # If model didn't include references, we inject the candidate pages.
        if "References:" not in out:
            pages_str = ", ".join(str(p) for p in sorted(set(ref_pages))) if ref_pages else "(none)"
            out = out.rstrip() + f"\nReferences: Page(s) {pages_str}\n"

        return out

    except Exception as e:
        print(f"Error with OpenAI API: {e}")
        return "An error occurred while processing the question."
    
def count_tokens(text, model="gpt-4o-mini"):
    import tiktoken
    try:
        enc = tiktoken.encoding_for_model(model)
    except KeyError:
        enc = tiktoken.get_encoding("cl100k_base")
    return len(enc.encode(text))


def split_text_by_tokens(text, max_tokens=200000, model="gpt-4o-mini"):
    import tiktoken
    try:
        enc = tiktoken.encoding_for_model(model)
    except KeyError:
        enc = tiktoken.get_encoding("cl100k_base")
    toks = enc.encode(text)
    chunks = []
    for i in range(0, len(toks), max_tokens):
        chunks.append(enc.decode(toks[i:i + max_tokens]))
    return chunks


def save_to_sharepoint_excel(email, filename, question, answer):
    access_token = session.get('access_token')
    if not access_token:
        print("No access token in session.")
        return

    headers = {
        'Authorization': f'Bearer {access_token}',
        'Content-Type': 'application/json'
    }

    site_name = "ProjectAI"
    drive_name = "Documents"
    file_path = "QnA_Storage.xlsx"
    table_name = "Table1"

    site_url = f"https://graph.microsoft.com/v1.0/sites/root:/sites/{site_name}"
    site_res = requests.get(site_url, headers=headers)
    site_id = site_res.json().get("id")

    drive_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drives"
    drive_res = requests.get(drive_url, headers=headers)
    drive_id = drive_res.json()['value'][0]['id']

    data = {
        "values": [[email, filename, question, answer, datetime.now().isoformat()]]
    }

    excel_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{file_path}:/workbook/tables/{table_name}/rows/add"
    res = requests.post(excel_url, headers=headers, json=data)

    if res.status_code == 201:
        print("Row added to Excel successfully.")
    else:
        print("Failed to add row to Excel:", res.text)


_OPENAI_SESSION = None


def _get_openai_session():
    global _OPENAI_SESSION
    if _OPENAI_SESSION is None:
        s = requests.Session()
        retry = Retry(
            total=5, read=5, connect=5,
            backoff_factor=0.6,
            status_forcelist=(429, 500, 502, 503, 504),
            allowed_methods=frozenset(["POST"]),
        )
        adapter = HTTPAdapter(max_retries=retry, pool_connections=10, pool_maxsize=10)
        s.mount("https://", adapter)
        s.mount("http://", adapter)
        _OPENAI_SESSION = s
    return _OPENAI_SESSION


_OPENAI_CLIENT = None


def _get_openai_client():
    global _OPENAI_CLIENT
    if _OPENAI_CLIENT is None:
        _OPENAI_CLIENT = OpenAI()
    return _OPENAI_CLIENT


def call_openai_text(
    model: str,
    prompt: str,
    temperature: float = 0.4,
    max_output_tokens: int | None = None,
    timeout_read: int = 480,
):
    """
    Robust OpenAI text call wrapper.

    - For gpt-5* uses Responses API.
    - Aggregates ALL output_text blocks (not just the first).
    - If response is incomplete (e.g., max_output_tokens), returns partial text
      and includes status + incomplete_reason in usage so callers can decide to continue.
    - Better retry/backoff on 429/5xx + respects Retry-After.
    """
    usage = {
        "input_tokens": None,
        "output_tokens": None,
        "status": None,
        "incomplete_reason": None,
        "response_id": None,
    }

    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY env var not set")

    # -----------------------------
    # Responses API branch (gpt-5*)
    # -----------------------------
    if model.startswith("gpt-5"):
        url = "https://api.openai.com/v1/responses"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        }

        payload = {
            "model": model,
            "input": prompt,
        }
        if max_output_tokens is not None:
            payload["max_output_tokens"] = int(max_output_tokens)

        session = _get_openai_session()

        # Retry strategy: network errors + 429 + common transient 5xx
        for attempt in range(1, 7):
            try:
                r = session.post(
                    url,
                    headers=headers,
                    json=payload,
                    timeout=(10, timeout_read),
                )

                # Handle throttling / transient server issues with retry/backoff
                if r.status_code in (429, 500, 502, 503, 504):
                    retry_after = r.headers.get("Retry-After")
                    if retry_after:
                        try:
                            sleep_s = float(retry_after)
                        except Exception:
                            sleep_s = None
                    else:
                        sleep_s = None

                    if sleep_s is None:
                        # exponential backoff with jitter
                        sleep_s = min(20.0, (2 ** (attempt - 1))) + random.uniform(0, 0.75)

                    print(f"[OpenAI] HTTP {r.status_code} (attempt {attempt}) -> sleeping {sleep_s:.1f}s")
                    time.sleep(sleep_s)
                    continue

                # Other errors: surface immediately (these are usually prompt/auth/etc.)
                if r.status_code >= 400:
                    raise RuntimeError(f"Error code: {r.status_code} - {r.text}")

                # Parse JSON safely
                try:
                    data = r.json()
                except Exception:
                    raw = r.content.decode("utf-8", errors="ignore")
                    data = parse_json_loose(raw) or {}

                # Capture metadata
                usage["status"] = data.get("status")
                usage["response_id"] = data.get("id")
                incomplete = data.get("incomplete_details") or {}
                usage["incomplete_reason"] = incomplete.get("reason")

                # ---- Extract ALL output_text chunks ----
                texts: list[str] = []

                # Preferred: structured output blocks
                for out in (data.get("output") or []):
                    for c in (out.get("content") or []):
                        if c.get("type") == "output_text":
                            t = c.get("text")
                            if isinstance(t, dict):
                                t = t.get("value", "")
                            if isinstance(t, str) and t:
                                texts.append(t)

                # Fallback: some SDKs place it here
                if not texts:
                    ot = data.get("output_text")
                    if isinstance(ot, str) and ot.strip():
                        texts.append(ot)

                text = "\n".join(t.strip("\n") for t in texts if isinstance(t, str)).strip()

                # Usage tokens
                u = data.get("usage") or {}
                usage["input_tokens"] = u.get("input_tokens")
                usage["output_tokens"] = u.get("output_tokens")

                # If we got ANY text, return it — even if incomplete
                if text:
                    return text, usage

                # No text at all: if incomplete due to max_output_tokens, still return empty but informative
                if usage["status"] == "incomplete":
                    # callers can decide what to do; but returning empty is safer than throwing here
                    return "", usage

                # Truly empty & not incomplete => error
                raise RuntimeError(
                    f"Responses API returned no output_text. status={usage['status']} "
                    f"incomplete_reason={usage['incomplete_reason']}"
                )

            except (requests.ReadTimeout, requests.ConnectionError,
                    requests.Timeout, requests.exceptions.ChunkedEncodingError) as e:
                sleep_s = min(20.0, (2 ** (attempt - 1))) + random.uniform(0, 0.75)
                print(f"[OpenAI] timeout/conn error (attempt {attempt}): {e}. Sleeping {sleep_s:.1f}s.")
                time.sleep(sleep_s)

        raise RuntimeError("OpenAI Responses call failed after retries")

    # -----------------------------
    # Chat Completions fallback
    # -----------------------------
    client = _get_openai_client().with_options(timeout=timeout_read)

    create_kwargs = {
        "model": model,
        "temperature": float(temperature),
        "messages": [{"role": "user", "content": prompt}],
    }
    if max_output_tokens is not None:
        create_kwargs["max_tokens"] = int(max_output_tokens)

    resp = client.chat.completions.create(**create_kwargs)
    text = (resp.choices[0].message.content or "").strip()

    try:
        u = resp.usage
        if u:
            usage["input_tokens"] = getattr(u, "prompt_tokens", None)
            usage["output_tokens"] = getattr(u, "completion_tokens", None)
    except Exception:
        pass

    return text, usage

@app.route('/admin/feedback')
def admin_feedback():
    if 'user' not in session or session['user']['email'] not in AUTHORIZED_EMAILS:
        return redirect(url_for('index'))
    feedback_list = Feedback.query.all()
    return render_template('admin_feedback.html', feedback_list=feedback_list)


@app.route("/ocr_progress/<job_id>")
def ocr_progress(job_id):
    def event_stream():
        while True:
            with OCR_JOBS_LOCK:
                job = OCR_JOBS.get(job_id)

            if not job:
                yield f"event: error\ndata: {json.dumps({'message': 'job not found'})}\n\n"
                return

            payload = {
                "status": job.get("status"),
                "message": job.get("message", ""),
                "done": job.get("done", 0),
                "total": job.get("total", 0),
                "page": job.get("page", 0),
                "chars": job.get("chars", 0),
                "last_conf": job.get("last_conf", None),
                "filename": job.get("filename"),
                "text_file": job.get("text_filepath", None),
            }

            yield f"data: {json.dumps(payload)}\n\n"

            if payload["status"] in ("done", "error", "cancelled"):
                return

            time.sleep(0.8)

    return Response(event_stream(), content_type="text/event-stream")


@app.route('/history')
def history():
    if 'user' not in session:
        return redirect(url_for('login'))

    user_email = session['user']['email']
    filename = request.args.get('filename', None)

    if filename:
        conversations = Conversation.query.filter_by(user_email=user_email, filename=filename).order_by(Conversation.timestamp.desc()).all()
    else:
        conversations = Conversation.query.filter_by(user_email=user_email).order_by(Conversation.timestamp.desc()).all()

    filenames = db.session.query(Conversation.filename).filter_by(user_email=user_email).distinct().all()
    return render_template('history.html', conversations=conversations, filenames=[f[0] for f in filenames])


@app.route('/admin/conversations')
def admin_conversations():
    if 'user' not in session or session['user']['email'] not in AUTHORIZED_EMAILS:
        return redirect(url_for('index'))
    all_conversations = Conversation.query.order_by(Conversation.timestamp.desc()).all()
    return render_template('admin_conversations.html', conversations=all_conversations)


def generate_answers_stream(filename):
    with open(f"temp_data/{filename}.txt", "r", encoding="utf-8") as f:
        extracted_text = f.read()

    questions = [
        "What is scope of this tender project?",
        "Where is the Work location?",
        "What is the closing date and time?"
    ]

    for question in questions:
        answer = get_answer_from_openai(question, extracted_text)
        time.sleep(1)
        yield f"data: {json.dumps({'question': question, 'answer': answer})}\n\n"


@app.route('/generate_answers_stream')
def generate_answer():
    filename = request.args.get('filename')
    if not filename or not os.path.exists(f"uploads/{filename}"):
        return Response("event: error\ndata: {\"answer\": \"PDF not found.\"}\n\n", mimetype='text/event-stream')

    text_filepath = os.path.join("temp_data", f"{filename}.txt")
    if not os.path.exists(text_filepath):
        return Response("event: error\ndata: {\"answer\": \"Text file not found.\"}\n\n", mimetype='text/event-stream')

    return Response(generate_answers_stream(filename), content_type='text/event-stream')


@app.route("/logout")
def logout():
    session.pop('email', None)
    return redirect(url_for('index'))


def create_app():
    app.secret_key = "your_secrete_key"
    app.config['UPLOAD_FOLDER'] = 'uploads'
    app.config["SQLALCHEMY_DATABASE_URI"] = os.getenv('DATABASE_URL')
    app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False

    db.init_app(app)
    migrate.init_app(app, db)
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    return app


def build_methodology_prompt_ecsa_iso(extracted_text, extra_context=""):
    context = (extracted_text or "")[:180000]
    style_pack = build_style_and_fewshot()

    style_guide = style_pack.get("style_guide", "")[:12000]
    mapping = style_pack.get("example_mapping", {})
    tender_snip = mapping.get("tender_snippet", "")[:4000]
    output_snip = mapping.get("output_snippet", "")[:6000]

    brief = (
        "You are a senior bid manager and project engineer.\n"
        "Goal: Generate a personalised, non-regurgitated Methodology & Preliminary Project Plan tailored to THIS RFP.\n\n"
        "STRICT STRUCTURE & RETURN FORMAT (VALID JSON ONLY):\n"
        "{\n"
        "  \"title\": string,\n"
        "  \"context_objectives\": string,\n"
        "  \"methodology\": [{\"phase\": string, \"steps\": [string]}],\n"
        "  \"plan_table\": [{\"work_package\": string, \"owner\": string, \"start\": string, \"finish\": string, \"duration_weeks\": number, \"dependencies\": string, \"deliverables\": string, \"acceptance_criteria\": string}],\n"
        "  \"project_management\": {\"governance\": string, \"controls\": string},\n"
        "  \"risk\": {\"framework\": string, \"top_risks\": [{\"name\": string, \"treatment\": string}]},\n"
        "  \"quality\": {\"qa_plan\": string, \"controls\": string},\n"
        "  \"additional_services\": [string],\n"
        "  \"assumptions\": [string],\n"
        "  \"references\": [string]\n"
        "}\n\n"
        "STYLE & LAYOUT: imitate the voice, density, headings, bullets, and flow from the STYLE GUIDE below.\n"
        "Avoid copying wording; instead, emulate tone and structure.\n"
        "ECSA, ISO 31000:2018, ISO 9001:2015 anchors must be explicit where relevant.\n"
    )
    if extra_context:
        brief += "\nAdditional preferences from user:\n" + extra_context

    fewshot = (
        "\n=== STYLE GUIDE (derived from our approved example DOCX) ===\n"
        + style_guide +
        "\n\n=== FEW-SHOT MAPPING (how tender text transforms to final prose) ===\n"
        "SOURCE_TENDER_SNIPPET:\n" + tender_snip +
        "\n\nEXPECTED_OUTPUT_STYLE_SNIPPET:\n" + output_snip + "\n"
    )

    return f"{brief}\n\n=== THIS RFP (source extract) ===\n{context}\n{fewshot}"


def _is_effectively_empty(value):
    if value is None:
        return True
    if isinstance(value, str):
        return not value.strip()
    if isinstance(value, (list, tuple)):
        if not value:
            return True
        return all(_is_effectively_empty(v) for v in value)
    if isinstance(value, dict):
        if not value:
            return True
        return all(_is_effectively_empty(v) for v in value.values())
    return False


def generate_methodology_json_single_call(model: str, full_text: str, extra_context: str):
    prompt = build_methodology_prompt_ecsa_iso(full_text, extra_context)

    try:
        text, _ = call_openai_text(
            model,
            prompt,
            max_output_tokens=12000,
            timeout_read=600,
        )
    except Exception as e:
        print("OpenAI single-call error:", e)
        raise

    if not isinstance(text, str) or not text.strip():
        print("[WARN] Empty response from model; using defaults.")
        return _min_default_doc()

    raw = parse_json_loose(text) or {}
    base = _min_default_doc()

    if isinstance(raw, list) and raw and isinstance(raw[0], dict):
        obj = raw[0]
    elif isinstance(raw, dict):
        obj = raw
    else:
        print("[WARN] Could not parse JSON from model; head:", str(text)[:400])
        return base

    def merged_field(key):
        v = obj.get(key)
        if _is_effectively_empty(v):
            return base[key]
        return v

    def merged_nested_dict(key, subkeys):
        raw_sub = obj.get(key)
        if not isinstance(raw_sub, dict):
            return base[key]
        out = {}
        for sk in subkeys:
            v = raw_sub.get(sk)
            if _is_effectively_empty(v):
                out[sk] = base[key][sk]
            else:
                out[sk] = v
        return out

    merged = {
        "title": merged_field("title"),
        "context_objectives": merged_field("context_objectives"),
        "methodology": merged_field("methodology"),
        "plan_table": merged_field("plan_table"),
        "project_management": merged_nested_dict("project_management", ["governance", "controls"]),
        "risk": merged_nested_dict("risk", ["framework", "top_risks"]),
        "quality": merged_nested_dict("quality", ["qa_plan", "controls"]),
        "additional_services": merged_field("additional_services"),
        "assumptions": merged_field("assumptions"),
        "references": merged_field("references"),

        # extra fields fall back to defaults
        "deliverables_register": merged_field("deliverables_register") if "deliverables_register" in obj else base["deliverables_register"],
        "document_control": merged_nested_dict("document_control", ["document_metadata", "review_approval"]) if "document_control" in obj else base["document_control"],
        "stakeholder_engagement": merged_nested_dict("stakeholder_engagement", ["approach", "channels"]) if "stakeholder_engagement" in obj else base["stakeholder_engagement"],
        "procurement_logic": merged_nested_dict("procurement_logic", ["stage_gates", "fee_repricing"]) if "procurement_logic" in obj else base["procurement_logic"],
        "value_engineering": merged_field("value_engineering") if "value_engineering" in obj else base["value_engineering"],
    }

    return merged


@app.route("/generate_methodology_plan", methods=["POST"])
def generate_methodology_plan():
    data = request.json or {}
    filename = data.get("filename")
    extra_context = data.get("preferences", "")

    if not filename:
        return jsonify({"message": "Filename not provided."}), 400

    text_path = os.path.join("temp_data", f"{filename}.txt")
    if not os.path.exists(text_path):
        return jsonify({"message": "Text file not found. Please upload/merge first."}), 404

    with open(text_path, "r", encoding="utf-8") as f:
        extracted_text = f.read()

    model = os.getenv("OPENAI_MODEL", "gpt-5-pro")

    doc_json = generate_methodology_json_by_sections(model, extracted_text, extra_context)

    raw_base = os.path.splitext(os.path.basename(filename))[0]
    safe_base = secure_filename(raw_base) or "rfp"
    MAX_BASENAME_LEN = 80
    if len(safe_base) > MAX_BASENAME_LEN:
        safe_base = safe_base[:MAX_BASENAME_LEN]

    pdf_name = f"{safe_base}_Methodology_Pictogram.pdf"
    pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], pdf_name)

    png_name = f"{safe_base}_Methodology_Pictogram_Page1.png"
    png_path = os.path.join(app.config['UPLOAD_FOLDER'], png_name)

    try:
        create_methodology_pictogram(doc_json, out_pdf_path=pdf_path, out_png_page1_path=png_path)
    except Exception as e:
        print("[WARN] Could not generate pictogram:", e)
        pdf_name = None
        png_name = None

    try:
        print("=== DOC_JSON (HEAD) ===")
        print(json.dumps(doc_json, indent=2)[:1500])
        print("=== END DOC_JSON ===")
    except Exception as e:
        print("Error printing doc_json:", e)

    docx_name = f"{safe_base}_Methodology_Plan_BRANDED.docx"
    docx_path = os.path.join(app.config['UPLOAD_FOLDER'], docx_name)

    try:
        render_into_template_after_toc(doc_json, COMPANY_TEMPLATE_PATH, docx_path)
    except Exception as e:
        print("DOCX render error:", e)
        return jsonify({"message": f"Could not render the branded DOCX: {e}"}), 500

    if not os.path.exists(docx_path):
        return jsonify({"message": "Document generation failed (file not found after save)."}), 500

    try:
        open(docx_path, "rb").close()
    except Exception as e:
        print("DOCX permission/read error:", e)
        return jsonify({"message": "Document generated but is not readable by the server."}), 500

    resp = {
        "message": "Generated successfully.",
        "structured": doc_json,
        "docx_download_url": url_for('uploaded_file', filename=docx_name),
    }

    if pdf_name:
        resp["pictogram_pdf_download_url"] = url_for('uploaded_file', filename=pdf_name)

    if png_name:
        resp["pictogram_png_download_url"] = url_for('uploaded_file', filename=png_name)

    return jsonify(resp)

# ===================== DUE DILIGENCE (WEB SEARCH) =====================
# Requires OpenAI Responses API tool: web_search
# You already have requests + retry session helpers.

def call_openai_with_web_search(
    model: str,
    prompt: str,
    max_output_tokens: int = 1200,
    timeout_read: int = 120,
    web_search_context_size: str = "medium",  # "low" | "medium" | "high"
):
    """
    Calls OpenAI Responses API with the web_search tool enabled.
    Returns (text, raw_response_json).
    """
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY env var not set")

    url = "https://api.openai.com/v1/responses"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }

    payload = {
        "model": model,
        "input": prompt,
        "max_output_tokens": int(max_output_tokens),
        "tools": [
            {
                "type": "web_search",
                "search_context_size": web_search_context_size
            }
        ],
        # Let the model decide tool usage:
        "tool_choice": "auto",
    }

    session = _get_openai_session()
    r = session.post(url, headers=headers, json=payload, timeout=(10, timeout_read))
    if r.status_code >= 400:
        raise RuntimeError(f"Error code: {r.status_code} - {r.text}")

    data = r.json()

    # Extract output text (same style as your call_openai_text)
    text = ""
    for out in (data.get("output") or []):
        content = out.get("content") or []
        for c in content:
            if c.get("type") == "output_text":
                t = c.get("text")
                if isinstance(t, dict):
                    t = t.get("value", "")
                if isinstance(t, str) and t.strip():
                    text = t.strip()
                    break
        if text:
            break

    if not text:
        text = (data.get("output_text") or "").strip()

    return text, data


def _dd_queries_phase1_identity(company: str, jurisdiction: str, industry: str, people: list[dict]) -> list[str]:
    q = []
    # Company identity / basics
    q += [
        f'{company} official website',
        f'{company} leadership team directors',
        f'{company} "About us" {jurisdiction}'.strip(),
    ]
    if industry:
        q.append(f'{company} {industry} profile')

    # People identity / role confirmation
    for p in (people or []):
        name = (p.get("name") or "").strip()
        role = (p.get("role") or "").strip()
        if not name:
            continue
        q += [
            f'"{name}" "{company}" {role}'.strip(),
            f'"{name}" LinkedIn "{company}"'.strip(),
            f'"{name}" "{company}" profile'.strip(),
        ]
    return list(dict.fromkeys([x for x in q if x]))


def _dd_queries_phase2_risk(company: str, jurisdiction: str, people: list[dict]) -> list[str]:
    # Company risk
    q = [
        f'"{company}" sanctions',
        f'"{company}" debarred OR blacklist OR procurement ban',
        f'"{company}" fraud OR corruption OR bribery OR "money laundering"',
        f'"{company}" regulator warning OR enforcement OR fine',
        f'"{company}" lawsuit OR litigation OR court case',
        f'"{company}" insolvency OR liquidation OR business rescue',
    ]
    if jurisdiction:
        q.append(f'"{company}" {jurisdiction} regulator enforcement')

    # Individual risk
    for p in (people or []):
        name = (p.get("name") or "").strip()
        if not name:
            continue
        q += [
            f'"{name}" sanctions OR debarred OR disqualified',
            f'"{name}" fraud OR corruption OR bribery OR "money laundering"',
            f'"{name}" investigation OR charged OR convicted',
            f'"{name}" PEP OR "politically exposed"',
            f'"{name}" lawsuit OR court case',
        ]
    return list(dict.fromkeys([x for x in q if x]))


def _dd_identity_badge_from_evidence(company_hits: int, people_hits: dict) -> tuple[str, str]:
    """
    Very simple and defensible scoring:
    - Confirmed: company has >=2 strong identity hits and >=1 hit for each person (or person list empty)
    - Limited: some hits but incomplete
    - Unverified: basically no identity hits
    """
    people = list(people_hits.keys())
    if company_hits >= 2 and (not people or all(people_hits.get(n, 0) >= 1 for n in people)):
        return "Identity confirmed", "Company and individuals have corroborating public identity/role signals."
    if company_hits >= 1 or any(v >= 1 for v in people_hits.values()):
        return "Limited public information", "Some public identity signals found, but coverage is incomplete or thin."
    return "Unverified", "Insufficient reliable public information to confirm identity/roles."


def _dd_build_synthesis_prompt(system_prompt: str, inputs: dict, evidence: dict) -> str:
    # Keep this compact and structured — the model will write the report using evidence only.
    return (
        system_prompt.strip()
        + "\n\nINPUTS:\n"
        + json.dumps(inputs, ensure_ascii=False, indent=2)
        + "\n\nEVIDENCE (PUBLIC SEARCH RESULTS, SUMMARIZED):\n"
        + json.dumps(evidence, ensure_ascii=False, indent=2)
        + "\n\nWRITE THE REPORT NOW using the exact required format."
    )


# Drop-in system prompt (from your earlier requirement)
DUE_DILIGENCE_SYSTEM_PROMPT = """
You are an AI assistant supporting iX’s Bid and Tender due diligence process.

GOAL
Conduct a high-level, public-information due diligence assessment of:
1) the customer company, and
2) the named key individuals involved in the deal.

Your purpose is to identify potential red flags relevant to fraud, corruption, sanctions, ethics, or major reputational risk.
You are NOT providing legal, compliance, or financial advice.
You are highlighting public signals for human review.

IMPORTANT PRINCIPLES
- Use ONLY lawful, publicly available information.
- Do NOT claim access to confidential, paid, private, or internal databases.
- Do NOT speculate or infer wrongdoing from absence of information.
- Clearly distinguish between:
  - Confirmed facts,
  - Credible allegations (with sources),
  - Unverified or speculative claims (generally exclude these).

SOURCE PRIORITY (CRITICAL)
You must assess information in TWO PHASES:

PHASE 1 — Identity & Role Confirmation (LOW RISK)
Permitted sources include:
- The company’s official website (About, Leadership, Directors, Media pages),
- Public professional profiles (e.g., LinkedIn),
- Public company registries and press releases.

This phase is used ONLY to:
- Confirm identity,
- Confirm role, seniority, and professional background,
- Establish whether the individual plausibly exists in the stated capacity.

Absence of adverse information in Phase 1 MUST NOT be treated as a red flag.

PHASE 2 — Risk & Adverse Signals (HIGHER SCRUTINY)
Permitted sources include:
- Reputable news outlets,
- Regulator announcements,
- Court judgments,
- Sanctions and watchlists,
- Enforcement or debarment lists,
- Public ESG or labour findings.

Only this phase may be used to identify red flags.

CONFIDENCE & LIMITATIONS
- If information is limited, state this explicitly.
- If no adverse information is found, say so clearly.
- Do NOT imply that lack of data equals risk.

OUTPUT REQUIREMENTS
Structure the response EXACTLY as follows:

1. Summary Assessment
2. Customer Company — 10 Due Diligence Questions
3. Key Individuals — 10 Due Diligence Questions per Person
4. Confidence Badge (Identity confidence level)

Always include:
“This is not legal or compliance advice; it is a public-information screening to support human decision-making.”
""".strip()

@app.route("/due-diligence_page", methods=["GET"])
def due_diligence_page():
    # Renders a standalone page for the feature
    return render_template("due_diligence.html")


@app.route("/due_diligence", methods=["POST"])
def due_diligence():
    """
    Two-phase public due diligence:
    Phase 1: identity confirmation searches
    Phase 2: risk/adverse searches
    Then synthesize report + badge.
    """
    data = request.json or {}
    company = (data.get("company") or "").strip()
    jurisdiction = (data.get("jurisdiction") or "").strip()
    industry = (data.get("industry") or "").strip()
    people = data.get("people") or []

    if not company:
        return jsonify({"message": "Customer company name is required."}), 400

    # ---- Models ----
    search_model = os.getenv("OPENAI_DD_SEARCH_MODEL", "gpt-4o-mini")
    analysis_model = os.getenv("OPENAI_DD_ANALYSIS_MODEL", "gpt-5-pro")

    # ---- Phase 1: Identity ----
    phase1_queries = _dd_queries_phase1_identity(company, jurisdiction, industry, people)

    # We’ll store a compact, UI-friendly evidence object
    evidence = {
        "inputs": {"company": company, "jurisdiction": jurisdiction, "industry": industry, "people": people},
        "phase1_identity": [],
        "phase2_risk": [],
    }

    # Track “hits” counts for the badge (simple signal: non-empty summaries)
    company_identity_hits = 0
    people_identity_hits = { (p.get("name") or "").strip(): 0 for p in people if (p.get("name") or "").strip() }

    try:
        for q in phase1_queries[:12]:  # cap for cost/time
            prompt = (
                "PHASE 1 (IDENTITY). Use web_search if helpful.\n"
                "Goal: find public info confirming identity/role only.\n"
                f"Query: {q}\n\n"
                "Return 3–6 bullet findings, each with: source name + date (if visible) + what it confirms.\n"
                "If nothing reliable is found, say: 'No reliable public information found.'"
            )
            txt, raw = call_openai_with_web_search(
                model=search_model,
                prompt=prompt,
                max_output_tokens=400,
                timeout_read=90,
                web_search_context_size="medium",
            )

            item = {"query": q, "summary": txt[:2000]}
            evidence["phase1_identity"].append(item)

            # crude hit counting (good enough for badge)
            if "No reliable public information found" not in (txt or "") and len((txt or "").strip()) > 40:
                if company.lower() in q.lower():
                    company_identity_hits += 1
                for person_name in list(people_identity_hits.keys()):
                    if person_name.lower() in q.lower():
                        people_identity_hits[person_name] += 1

        badge, badge_reason = _dd_identity_badge_from_evidence(company_identity_hits, people_identity_hits)

        # ---- Phase 2: Risk ----
        phase2_queries = _dd_queries_phase2_risk(company, jurisdiction, people)

        for q in phase2_queries[:16]:  # cap
            prompt = (
                "PHASE 2 (RISK). Use web_search if helpful.\n"
                "Goal: find credible adverse signals (sanctions, enforcement, corruption, fraud, debarment, serious litigation).\n"
                f"Query: {q}\n\n"
                "Return up to 6 bullets. Each bullet must include: (a) what happened, (b) date, (c) source.\n"
                "If nothing reliable is found, say: 'No reliable adverse public information found.'"
            )
            txt, raw = call_openai_with_web_search(
                model=search_model,
                prompt=prompt,
                max_output_tokens=450,
                timeout_read=90,
                web_search_context_size="high",
            )
            evidence["phase2_risk"].append({"query": q, "summary": txt[:2200]})

        # ---- Synthesis (NO web tool; uses evidence only) ----
        inputs_obj = {
            "CustomerCompanyName": company,
            "Jurisdiction": jurisdiction or "Not specified",
            "Industry": industry or "Not specified",
            "KeyIndividuals": people,
        }

        synthesis_prompt = _dd_build_synthesis_prompt(DUE_DILIGENCE_SYSTEM_PROMPT, inputs_obj, evidence)

        report_text, usage = call_openai_text(
            model=analysis_model,
            prompt=synthesis_prompt,
            temperature=0.2,
            max_output_tokens=5500,
            timeout_read=480,
        )

        # Append badge section if the model didn’t include it cleanly
        # (We still return badge separately for UI)
        if "Confidence Badge" not in (report_text or ""):
            report_text = (report_text or "").strip() + (
                f"\n\n4. Confidence Badge (Identity confidence level)\n"
                f"- Badge: {badge}\n"
                f"- Rationale: {badge_reason}\n"
                f"- Note: This is not legal or compliance advice; it is a public-information screening to support human decision-making.\n"
            )

        return jsonify({
            "report": report_text,
            "evidence": evidence,
            "identity_badge": badge,
            "identity_badge_reason": badge_reason,
            "models": {
                "search_model": search_model,
                "analysis_model": analysis_model
            }
        })

    except Exception as e:
        print("[DUE_DILIGENCE] error:", e)
        return jsonify({
            "message": "Due diligence failed (search or synthesis error).",
            "error": str(e)
        }), 500

if __name__ == '__main__':
    app.run(debug=True)